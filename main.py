#!/usr/bin/env python3
"""
🤖 БОТ "ТРИТИКА" (ТЕНДЕРПОИСК)
Интеллектуальный ассистент для поиска тендеров
"""

import os
import asyncio
import logging
import sqlite3
import tempfile
import json
import io
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path

from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    ReplyKeyboardMarkup, KeyboardButton,
    InlineKeyboardMarkup, InlineKeyboardButton,
    ReplyKeyboardRemove, BufferedInputFile, FSInputFile
)
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from docx import Document
from docx.shared import Inches

# Импорты для HTTP сервера Railway
import aiohttp
from aiohttp import web

# =========== НАСТРОЙКИ ===========
BOT_TOKEN = os.getenv("BOT_TOKEN", "")
ADMIN_ID = int(os.getenv("ADMIN_ID", "0")) if os.getenv("ADMIN_ID") else None
PORT = int(os.getenv("PORT", 8080))

# Настройки времени работы (пн-чт 8:30-17:30 пт 8:30-16:30)
WORK_START_HOUR = 9
WORK_END_HOUR = 17
WORK_DAYS = [0, 1, 2, 3, 4]  # Пн-Пт

# Ссылка на файл анкеты в GitHub
ANKETA_GITHUB_URL = "https://github.com/Oleg331987/rassilka/raw/main/Anketa.docx"
ANKETA_LOCAL_PATH = "Anketa.docx"

# Папка для временных файлов выгрузок
EXPORTS_DIR = "exports"
os.makedirs(EXPORTS_DIR, exist_ok=True)

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

print("="*60)
print("🤖 ЗАГРУЗКА БОТА ТРИТИКА (ТЕНДЕРПОИСК)")
print("="*60)

# =========== ИНИЦИАЛИЗАЦИЯ БОТА ===========
try:
    bot = Bot(
        token=BOT_TOKEN,
        default=DefaultBotProperties(parse_mode=ParseMode.HTML)
    )
    storage = MemoryStorage()
    dp = Dispatcher(storage=storage)
    logger.info("✅ Бот инициализирован")
except Exception as e:
    logger.error(f"❌ Ошибка инициализации бота: {e}")
    exit(1)

# =========== ФУНКЦИЯ ДЛЯ СОЗДАНИЯ ЗАПОЛНЕННОЙ АНКЕТЫ ===========
def create_filled_anketa(user_data: dict) -> Optional[str]:
    """Создание заполненной анкеты на основе данных пользователя"""
    try:
        # Создаем новый документ
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('Анкета для поиска тендеров', 0)
        title.alignment = 1
        
        # Информация о заполнении
        doc.add_paragraph(f'Дата заполнения: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('Заполнено через бота Тритика')
        
        # Информация о компании
        doc.add_heading('Информация о компании', level=1)
        
        # Заполняем поля (новый порядок)
        fields = [
            ('1. Сфера деятельности компании:', user_data.get('activity', 'Не указано')),
            ('2. Регионы работы (города, области):', user_data.get('region', 'Не указано')),
            ('3. Предпочтительный бюджет контрактов:', user_data.get('budget', 'Не указано')),
            ('4. Ключевые слова для поиска (через запятую):', user_data.get('keywords', 'Не указано')),
            ('5. Название компании:', user_data.get('company_name', 'Не указано')),
            ('6. ФИО полностью:', user_data.get('full_name', 'Не указано')),
            ('7. Телефон для связи:', user_data.get('phone', 'Не указано')),
            ('8. Email для отправки тендеров:', user_data.get('email', 'Не указано')),
        ]
        
        for label, value in fields:
            p = doc.add_paragraph()
            p.add_run(label).bold = True
            doc.add_paragraph(value)
            doc.add_paragraph()  # Пустая строка
        
        # Подвал
        doc.add_page_break()
        doc.add_paragraph('\n\n')
        doc.add_paragraph('Анкета заполена через Telegram-бота Тритика')
        doc.add_paragraph('https://t.me/tritika_tender_bot')
        
        # Сохраняем во временный файл
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_path = temp_file.name
        doc.save(temp_path)
        temp_file.close()
        
        logger.info(f"✅ Файл анкеты создан: {temp_path}")
        return temp_path
        
    except Exception as e:
        logger.error(f"❌ Ошибка создания заполненной анкеты: {e}")
        return None

# =========== СКАЧИВАНИЕ ФАЙЛА ANKETA.DOCX ===========
async def download_anketa_file():
    """Скачивание файла анкеты с GitHub"""
    try:
        print("⬇️ Скачиваю файл анкеты с GitHub...")
        async with aiohttp.ClientSession() as session:
            async with session.get(ANKETA_GITHUB_URL, timeout=30) as response:
                if response.status == 200:
                    content = await response.read()
                    # Создаем папку, если её нет
                    os.makedirs(os.path.dirname(ANKETA_LOCAL_PATH), exist_ok=True)
                    with open(ANKETA_LOCAL_PATH, 'wb') as f:
                        f.write(content)
                    print(f"✅ Файл анкеты сохранен: {ANKETA_LOCAL_PATH} ({len(content)} байт)")
                    return True
                else:
                    print(f"❌ Ошибка скачивания файла: HTTP {response.status}")
                    return False
    except Exception as e:
        print(f"❌ Ошибка скачивания анкеты: {e}")
        return False

# =========== БАЗА ДАННЫХ ===========
class Database:
    def __init__(self, db_name="tenders.db"):
        self.db_name = db_name
        self.init_db()
    
    def init_db(self):
        """Инициализация базы данных с новыми таблицами"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        # Пользователи - добавляем поле для управления рассылкой
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER UNIQUE,
            username TEXT,
            first_name TEXT,
            last_name TEXT,
            phone TEXT,
            email TEXT,
            company TEXT,
            activity TEXT,
            region TEXT,
            is_active BOOLEAN DEFAULT 1,
            has_filled_questionnaire BOOLEAN DEFAULT 0,
            mailing_subscribed BOOLEAN DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            last_mailing_date TIMESTAMP
        )
        ''')
        
        # Анкеты (отдельная таблица для истории) с статусом
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS questionnaires (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            full_name TEXT,
            company_name TEXT,
            phone TEXT,
            email TEXT,
            activity TEXT,
            region TEXT,
            budget TEXT,
            keywords TEXT,
            filled_anketa_path TEXT,
            status TEXT DEFAULT 'new',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        ''')
        
        # Выгрузки тендеров - УПРОЩЕННАЯ ВЕРСИЯ
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS tender_exports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            file_path TEXT,
            file_name TEXT,
            sent_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            sent_by TEXT DEFAULT 'bot',
            status TEXT DEFAULT 'pending',
            admin_notified BOOLEAN DEFAULT 0,
            follow_up_sent BOOLEAN DEFAULT 0,
            follow_up_at TIMESTAMP,
            follow_up_response TEXT,
            follow_up_scheduled BOOLEAN DEFAULT 0
        )
        ''')
        
        # Рассылки (ручные) - основная таблица
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS manual_mailings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            admin_id INTEGER,
            mailing_text TEXT,
            mailing_type TEXT,
            filter_criteria TEXT,
            sent_count INTEGER DEFAULT 0,
            failed_count INTEGER DEFAULT 0,
            feedback_count INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            sent_at TIMESTAMP
        )
        ''')
        
        # Отправленные сообщения рассылки (каждому пользователю)
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS sent_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            mailing_id INTEGER,
            user_id INTEGER,
            telegram_message_id INTEGER,
            sent_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            feedback_received BOOLEAN DEFAULT 0
        )
        ''')
        
        # Обратная связь по рассылкам
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS mailing_feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            mailing_id INTEGER,
            user_id INTEGER,
            sent_message_id INTEGER,
            feedback_type TEXT,
            feedback_text TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        ''')
        
        # Сообщения менеджеру
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS manager_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            message_type TEXT,
            message_text TEXT,
            file_id TEXT,
            file_name TEXT,
            admin_notified BOOLEAN DEFAULT 0,
            processed BOOLEAN DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        ''')
        
        # Запросы контактов для выгрузок
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS contact_requests (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            export_id INTEGER,
            requested_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            completed BOOLEAN DEFAULT 0,
            completed_at TIMESTAMP
        )
        ''')
        
        conn.commit()
        conn.close()
        logger.info("✅ База данных инициализирована")
    
    def add_user(self, user_id: int, username: str, first_name: str, last_name: str = ""):
        """Добавление пользователя"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT OR IGNORE INTO users (user_id, username, first_name, last_name)
        VALUES (?, ?, ?, ?)
        ''', (user_id, username, first_name, last_name))
        
        conn.commit()
        conn.close()
        return True
    
    def save_questionnaire_partial(self, user_id: int, data: dict):
        """Сохранение частичной анкеты (только вопросы 1-4)"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT INTO questionnaires 
        (user_id, activity, region, budget, keywords, status)
        VALUES (?, ?, ?, ?, ?, 'partial')
        ''', (
            user_id,
            data.get('activity'),
            data.get('region'),
            data.get('budget'),
            data.get('keywords')
        ))
        
        conn.commit()
        last_id = cursor.lastrowid
        
        cursor.execute('''
        UPDATE users 
        SET activity = ?, region = ?, has_filled_questionnaire = 1
        WHERE user_id = ?
        ''', (
            data.get('activity'),
            data.get('region'),
            user_id
        ))
        
        conn.commit()
        conn.close()
        
        return last_id
    
    def save_questionnaire(self, user_id: int, data: dict, anketa_path: str = None):
        """Сохранение полной анкеты (все 8 вопросов)"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT INTO questionnaires 
        (user_id, full_name, company_name, phone, email, activity, region, budget, keywords, filled_anketa_path, status)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'complete')
        ''', (
            user_id,
            data.get('full_name'),
            data.get('company_name'),
            data.get('phone'),
            data.get('email'),
            data.get('activity'),
            data.get('region'),
            data.get('budget'),
            data.get('keywords'),
            anketa_path
        ))
        
        conn.commit()
        last_id = cursor.lastrowid
        
        cursor.execute('''
        UPDATE users 
        SET phone = ?, email = ?, company = ?, activity = ?, region = ?, has_filled_questionnaire = 1
        WHERE user_id = ?
        ''', (
            data.get('phone'),
            data.get('email'),
            data.get('company_name'),
            data.get('activity'),
            data.get('region'),
            user_id
        ))
        
        conn.commit()
        conn.close()
        
        return last_id
    
    def update_partial_to_complete(self, user_id: int, data: dict):
        """Обновление частичной анкеты до полной (добавление контактов)"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        UPDATE questionnaires 
        SET company_name = ?, full_name = ?, phone = ?, email = ?, status = 'complete'
        WHERE user_id = ? AND status = 'partial'
        ''', (
            data.get('company_name'),
            data.get('full_name'),
            data.get('phone'),
            data.get('email'),
            user_id
        ))
        
        cursor.execute('''
        UPDATE users 
        SET phone = ?, email = ?, company = ?
        WHERE user_id = ?
        ''', (
            data.get('phone'),
            data.get('email'),
            data.get('company_name'),
            user_id
        ))
        
        conn.commit()
        conn.close()
        return True
    
    def create_tender_export(self, user_id: int, file_path: str = None, file_name: str = None):
        """Создание записи о выгрузке тендеров - УПРОЩЕННАЯ ВЕРСИЯ"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT INTO tender_exports 
        (user_id, file_path, file_name, follow_up_scheduled)
        VALUES (?, ?, ?, ?)
        ''', (user_id, file_path, file_name, 1))
        
        conn.commit()
        export_id = cursor.lastrowid
        conn.close()
        
        return export_id
    
    def mark_export_completed(self, export_id: int, admin_name: str = "Олег"):
        """Отметка выполнения выгрузки администратором"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        UPDATE tender_exports 
        SET sent_by = ?, status = 'completed', admin_notified = 1
        WHERE id = ?
        ''', (admin_name, export_id))
        
        conn.commit()
        conn.close()
    
    def save_export_file(self, export_id: int, file_path: str, file_name: str):
        """Сохранение пути к файлу выгрузки"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        UPDATE tender_exports 
        SET file_path = ?, file_name = ?, status = 'pending'
        WHERE id = ?
        ''', (file_path, file_name, export_id))
        
        conn.commit()
        conn.close()
    
    def get_exports_for_followup(self):
        """Получение выгрузок, для которых нужно отправить follow-up"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        one_hour_ago = (datetime.now() - timedelta(hours=1)).strftime('%Y-%m-%d %H:%M:%S')
        
        cursor.execute('''
        SELECT te.*, u.username, u.first_name, u.last_name
        FROM tender_exports te
        JOIN users u ON te.user_id = u.user_id
        WHERE te.status = 'completed' 
        AND te.follow_up_scheduled = 1
        AND te.follow_up_sent = 0
        AND te.sent_at <= ?
        ''', (one_hour_ago,))
        
        exports = cursor.fetchall()
        conn.close()
        
        return exports
    
    def mark_followup_sent(self, export_id: int):
        """Отметка, что follow-up отправлен"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        UPDATE tender_exports 
        SET follow_up_sent = 1, follow_up_at = datetime('now')
        WHERE id = ?
        ''', (export_id,))
        
        conn.commit()
        conn.close()
    
    def save_followup_response(self, export_id: int, response: str):
        """Сохранение ответа на follow-up"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        UPDATE tender_exports 
        SET follow_up_response = ?
        WHERE id = ?
        ''', (response, export_id))
        
        conn.commit()
        conn.close()
    
    def toggle_user_mailing_subscription(self, user_id: int):
        """Включение/выключение подписки на рассылку"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('SELECT mailing_subscribed FROM users WHERE user_id = ?', (user_id,))
        current = cursor.fetchone()
        
        if current:
            new_status = not bool(current[0])
            cursor.execute('''
            UPDATE users 
            SET mailing_subscribed = ?
            WHERE user_id = ?
            ''', (1 if new_status else 0, user_id))
            
            conn.commit()
            conn.close()
            return new_status
        
        conn.close()
        return None
    
    def get_user_mailing_status(self, user_id: int):
        """Получение статуса подписки на рассылку"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT mailing_subscribed, username, first_name, last_name 
        FROM users 
        WHERE user_id = ?
        ''', (user_id,))
        
        result = cursor.fetchone()
        conn.close()
        
        if result:
            return {
                'subscribed': bool(result[0]),
                'username': result[1],
                'first_name': result[2],
                'last_name': result[3]
            }
        return None
    
    def get_users_by_filter(self, filter_type: str):
        """Получение пользователей по фильтру с учетом подписки"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        if filter_type == "all":
            cursor.execute('''
            SELECT user_id, username, first_name, last_name, company, mailing_subscribed
            FROM users 
            WHERE is_active = 1 AND mailing_subscribed = 1
            ''')
        elif filter_type == "with_questionnaire":
            cursor.execute('''
            SELECT user_id, username, first_name, last_name, company, mailing_subscribed
            FROM users 
            WHERE is_active = 1 AND has_filled_questionnaire = 1 AND mailing_subscribed = 1
            ''')
        elif filter_type == "without_questionnaire":
            cursor.execute('''
            SELECT user_id, username, first_name, last_name, company, mailing_subscribed
            FROM users 
            WHERE is_active = 1 AND has_filled_questionnaire = 0 AND mailing_subscribed = 1
            ''')
        elif filter_type == "recent_week":
            cursor.execute('''
            SELECT user_id, username, first_name, last_name, company, mailing_subscribed
            FROM users 
            WHERE is_active = 1 AND mailing_subscribed = 1 
            AND date(created_at) >= date('now', '-7 days')
            ''')
        elif filter_type == "subscribed":
            cursor.execute('''
            SELECT user_id, username, first_name, last_name, company, mailing_subscribed
            FROM users 
            WHERE is_active = 1 AND mailing_subscribed = 1
            ''')
        elif filter_type == "unsubscribed":
            cursor.execute('''
            SELECT user_id, username, first_name, last_name, company, mailing_subscribed
            FROM users 
            WHERE is_active = 1 AND mailing_subscribed = 0
            ''')
        else:
            conn.close()
            return []
        
        users = cursor.fetchall()
        conn.close()
        
        return users
    
    def get_all_users_with_subscription(self, limit: int = 50):
        """Получение всех пользователей с информацией о подписке"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT user_id, username, first_name, last_name, company, 
               mailing_subscribed, has_filled_questionnaire, created_at
        FROM users 
        WHERE is_active = 1
        ORDER BY created_at DESC
        LIMIT ?
        ''', (limit,))
        
        users = cursor.fetchall()
        conn.close()
        
        return users
    
    def create_manual_mailing(self, admin_id: int, mailing_text: str, mailing_type: str, filter_criteria: str):
        """Создание ручной рассылки"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT INTO manual_mailings (admin_id, mailing_text, mailing_type, filter_criteria)
        VALUES (?, ?, ?, ?)
        ''', (admin_id, mailing_text, mailing_type, filter_criteria))
        
        conn.commit()
        mailing_id = cursor.lastrowid
        conn.close()
        
        return mailing_id
    
    def save_sent_message(self, mailing_id: int, user_id: int, telegram_message_id: int):
        """Сохранение отправленного сообщения"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT INTO sent_messages (mailing_id, user_id, telegram_message_id)
        VALUES (?, ?, ?)
        ''', (mailing_id, user_id, telegram_message_id))
        
        conn.commit()
        message_id = cursor.lastrowid
        conn.close()
        
        return message_id
    
    def update_mailing_stats(self, mailing_id: int, sent_count: int, failed_count: int):
        """Обновление статистики рассылки"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        UPDATE manual_mailings 
        SET sent_count = ?, failed_count = ?, sent_at = datetime('now')
        WHERE id = ?
        ''', (sent_count, failed_count, mailing_id))
        
        conn.commit()
        conn.close()
    
    def save_mailing_feedback(self, mailing_id: int, user_id: int, sent_message_id: int, 
                             feedback_type: str, feedback_text: str = ""):
        """Сохранение обратной связи по рассылке"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT INTO mailing_feedback 
        (mailing_id, user_id, sent_message_id, feedback_type, feedback_text)
        VALUES (?, ?, ?, ?, ?)
        ''', (mailing_id, user_id, sent_message_id, feedback_type, feedback_text))
        
        cursor.execute('''
        UPDATE manual_mailings 
        SET feedback_count = feedback_count + 1
        WHERE id = ?
        ''', (mailing_id,))
        
        cursor.execute('''
        UPDATE sent_messages 
        SET feedback_received = 1
        WHERE id = ?
        ''', (sent_message_id,))
        
        conn.commit()
        feedback_id = cursor.lastrowid
        conn.close()
        
        return feedback_id
    
    def get_sent_message_by_telegram_id(self, user_id: int, telegram_message_id: int):
        """Получение отправленного сообщения по ID Telegram"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT sm.*, mm.mailing_text
        FROM sent_messages sm
        JOIN manual_mailings mm ON sm.mailing_id = mm.id
        WHERE sm.user_id = ? AND sm.telegram_message_id = ?
        ''', (user_id, telegram_message_id))
        
        result = cursor.fetchone()
        conn.close()
        
        return result
    
    def get_mailing_feedback(self, mailing_id: int):
        """Получение обратной связи по рассылке"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT mf.*, u.username, u.first_name, u.last_name
        FROM mailing_feedback mf
        JOIN users u ON mf.user_id = u.user_id
        WHERE mf.mailing_id = ?
        ORDER BY mf.created_at DESC
        ''', (mailing_id,))
        
        feedback = cursor.fetchall()
        conn.close()
        
        return feedback
    
    def get_mailing_feedback_for_user(self, user_id: int):
        """Получение обратной связи по пользователю"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT mf.*, mm.mailing_text
        FROM mailing_feedback mf
        JOIN manual_mailings mm ON mf.mailing_id = mm.id
        WHERE mf.user_id = ?
        ORDER BY mf.created_at DESC
        LIMIT 10
        ''', (user_id,))
        
        feedback = cursor.fetchall()
        conn.close()
        
        return feedback
    
    def save_manager_message(self, user_id: int, message_type: str, message_text: str, file_id: str = None, file_name: str = None):
        """Сохранение сообщения менеджеру"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT INTO manager_messages (user_id, message_type, message_text, file_id, file_name)
        VALUES (?, ?, ?, ?, ?)
        ''', (user_id, message_type, message_text, file_id, file_name))
        
        conn.commit()
        message_id = cursor.lastrowid
        conn.close()
        
        return message_id
    
    def get_pending_exports(self):
        """Получение ожидающих выгрузок"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT te.*, u.username, u.first_name, u.last_name, u.email, u.phone
        FROM tender_exports te
        JOIN users u ON te.user_id = u.user_id
        WHERE te.status = 'pending'
        ORDER BY te.sent_at DESC
        LIMIT 10
        ''')
        
        exports = cursor.fetchall()
        conn.close()
        
        return exports
    
    def get_user_by_id(self, user_id: int):
        """Получение пользователя по ID"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT * FROM users 
        WHERE user_id = ?
        ''', (user_id,))
        
        user = cursor.fetchone()
        conn.close()
        
        return user
    
    def get_export_by_id(self, export_id: int):
        """Получение выгрузки по ID"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT te.*, u.username, u.first_name, u.last_name, u.email, u.phone
        FROM tender_exports te
        JOIN users u ON te.user_id = u.user_id
        WHERE te.id = ?
        ''', (export_id,))
        
        export = cursor.fetchone()
        conn.close()
        
        return export
    
    def get_user_exports(self, user_id: int):
        """Получение всех выгрузок пользователя"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT te.*
        FROM tender_exports te
        WHERE te.user_id = ?
        ORDER BY te.sent_at DESC
        LIMIT 20
        ''', (user_id,))
        
        exports = cursor.fetchall()
        conn.close()
        
        return exports
    
    def get_statistics(self, days: int = 14):
        """Получение статистики за указанный период"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
        
        cursor.execute('''
        SELECT COUNT(*) as count FROM users 
        WHERE date(created_at) >= ?
        ''', (start_date,))
        new_users = cursor.fetchone()['count']
        
        cursor.execute('''
        SELECT COUNT(*) as count FROM tender_exports 
        WHERE date(sent_at) >= ? AND status = 'completed'
        ''', (start_date,))
        exports_completed = cursor.fetchone()['count']
        
        cursor.execute('''
        SELECT COUNT(*) as count FROM manager_messages 
        WHERE date(created_at) >= ?
        ''', (start_date,))
        manager_messages = cursor.fetchone()['count']
        
        cursor.execute('''
        SELECT 
            COUNT(*) as count, 
            SUM(sent_count) as total_sent,
            SUM(feedback_count) as total_feedback
        FROM manual_mailings 
        WHERE date(created_at) >= ?
        ''', (start_date,))
        mailings = cursor.fetchone()
        
        cursor.execute('''
        SELECT 
            SUM(CASE WHEN mailing_subscribed = 1 THEN 1 ELSE 0 END) as subscribed,
            SUM(CASE WHEN mailing_subscribed = 0 THEN 1 ELSE 0 END) as unsubscribed
        FROM users 
        WHERE is_active = 1
        ''')
        subscriptions = cursor.fetchone()
        
        cursor.execute('''
        SELECT COUNT(*) as count FROM questionnaires 
        WHERE date(created_at) >= ?
        ''', (start_date,))
        new_questionnaires = cursor.fetchone()['count']
        
        conn.close()
        
        return {
            'new_users': new_users,
            'exports_completed': exports_completed,
            'manager_messages': manager_messages,
            'mailings_count': mailings['count'] if mailings and mailings['count'] else 0,
            'mailings_sent': mailings['total_sent'] if mailings and mailings['total_sent'] else 0,
            'mailings_feedback': mailings['total_feedback'] if mailings and mailings['total_feedback'] else 0,
            'subscribed_users': subscriptions['subscribed'] if subscriptions and subscriptions['subscribed'] else 0,
            'unsubscribed_users': subscriptions['unsubscribed'] if subscriptions and subscriptions['unsubscribed'] else 0,
            'new_questionnaires': new_questionnaires
        }
    
    def is_working_hours(self):
        """Проверка рабочего времени"""
        now = datetime.now()
        
        if now.weekday() not in WORK_DAYS:
            return False
        
        if now.hour < WORK_START_HOUR or now.hour >= WORK_END_HOUR:
            return False
        
        return True
    
    def get_next_working_time(self):
        """Получение следующего рабочего времени"""
        now = datetime.now()
        
        if self.is_working_hours():
            return now
        
        days_to_add = 1
        while (now.weekday() + days_to_add) % 7 not in WORK_DAYS:
            days_to_add += 1
        
        next_work_day = now + timedelta(days=days_to_add)
        return next_work_day.replace(hour=WORK_START_HOUR, minute=0, second=0, microsecond=0)

    def get_partial_questionnaires(self):
        """Получение частичных анкет (только 1-4 вопросы)"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT q.*, u.username 
        FROM questionnaires q
        LEFT JOIN users u ON q.user_id = u.user_id
        WHERE q.status = 'partial'
        ORDER BY q.created_at DESC
        LIMIT 20
        ''')
        
        questionnaires = cursor.fetchall()
        conn.close()
        
        return questionnaires
    
    def get_complete_questionnaires(self):
        """Получение полных анкет"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT q.*, u.username 
        FROM questionnaires q
        LEFT JOIN users u ON q.user_id = u.user_id
        WHERE q.status = 'complete'
        ORDER BY q.created_at DESC
        LIMIT 20
        ''')
        
        questionnaires = cursor.fetchall()
        conn.close()
        
        return questionnaires
    
    def create_contact_request(self, user_id: int, export_id: int):
        """Создание запроса контактов для выгрузки"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT INTO contact_requests (user_id, export_id)
        VALUES (?, ?)
        ''', (user_id, export_id))
        
        conn.commit()
        request_id = cursor.lastrowid
        conn.close()
        
        return request_id
    
    def mark_contact_request_completed(self, export_id: int):
        """Отметка запроса контактов как выполненного"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        UPDATE contact_requests 
        SET completed = 1, completed_at = datetime('now')
        WHERE export_id = ?
        ''', (export_id,))
        
        conn.commit()
        conn.close()
    
    def create_tender_export_without_file(self, user_id: int):
        """Создание записи о выгрузке без файла"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        INSERT INTO tender_exports 
        (user_id, follow_up_scheduled, status)
        VALUES (?, ?, 'pending')
        ''', (user_id, 1))
        
        conn.commit()
        export_id = cursor.lastrowid
        conn.close()
        
        return export_id

    def has_complete_questionnaire(self, user_id: int):
        """Проверка, есть ли у пользователя полная анкета с контактами"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT COUNT(*) as count 
        FROM questionnaires 
        WHERE user_id = ? AND status = 'complete'
        AND full_name IS NOT NULL 
        AND phone IS NOT NULL 
        AND email IS NOT NULL
        ''', (user_id,))
        
        result = cursor.fetchone()
        conn.close()
        
        return result[0] > 0 if result else False

    def get_last_complete_questionnaire(self, user_id: int):
        """Получение последней полной анкеты пользователя"""
        conn = sqlite3.connect(self.db_name)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT * 
        FROM questionnaires 
        WHERE user_id = ? AND status = 'complete'
        ORDER BY created_at DESC
        LIMIT 1
        ''', (user_id,))
        
        questionnaire = cursor.fetchone()
        conn.close()
        
        return questionnaire

db = Database()

# =========== HTTP ОБРАБОТЧИКИ ДЛЯ RAILWAY ===========
async def health_check(request):
    """Health check endpoint для Railway"""
    return web.Response(text="OK", status=200)

async def status_check(request):
    """Статус бота"""
    try:
        bot_info = await bot.get_me()
        stats = db.get_statistics(7)
        
        return web.json_response({
            "status": "running",
            "bot": f"@{bot_info.username}",
            "name": bot_info.first_name,
            "statistics": stats,
            "timestamp": datetime.now().isoformat()
        })
    except Exception as e:
        return web.json_response({
            "status": "error",
            "error": str(e)
        }, status=500)

# =========== КЛАВИАТУРЫ ===========
def get_main_keyboard():
    """Главная клавиатура (УБРАНА КНОПКА ПОДЕЛИТЬСЯ ТЕЛЕФОНОМ)"""
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="📝 Заполнить анкету онлайн")],
            [KeyboardButton(text="📥 Скачать анкету в Word")],
            [KeyboardButton(text="📤 Написать менеджеру")],
            [KeyboardButton(text="📊 Мои выгрузки")],
            [KeyboardButton(text="📞 Контакты"), KeyboardButton(text="ℹ️ Помощь")]
        ],
        resize_keyboard=True,
        input_field_placeholder="Выберите действие..."
    )

def get_phone_keyboard():
    """Клавиатура для ввода телефона с кнопкой поделиться"""
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="📱 Поделиться телефоном", request_contact=True)],
            [KeyboardButton(text="📝 Ввести вручную")],
            [KeyboardButton(text="❌ Отмена")]
        ],
        resize_keyboard=True
    )

def get_phone_keyboard_simple():
    """Упрощенная клавиатура для ввода телефона (только кнопка поделиться и отмена)"""
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="📱 Поделиться телефоном", request_contact=True)],
            [KeyboardButton(text="❌ Отмена")]
        ],
        resize_keyboard=True
    )

def get_admin_keyboard():
    """Клавиатура администратора"""
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="📊 Частичные анкеты"), KeyboardButton(text="📤 Отправить выгрузку")],
            [KeyboardButton(text="📈 Статистика"), KeyboardButton(text="📨 Создать рассылку")],
            [KeyboardButton(text="👥 Управление подписками"), KeyboardButton(text="📩 Сообщения менеджеру")],
            [KeyboardButton(text="📋 Обратная связь"), KeyboardButton(text="⚙️ Настройки")],
            [KeyboardButton(text="👤 Режим пользователя")]
        ],
        resize_keyboard=True
    )

def get_cancel_keyboard():
    """Клавиатура отмены"""
    return ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="❌ Отмена")]],
        resize_keyboard=True
    )

def get_follow_up_keyboard(export_id: int):
    """Клавиатура для follow-up"""
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="✅ Да, нашел подходящее", callback_data=f"follow_yes_{export_id}"),
                InlineKeyboardButton(text="❌ Нет, не нашел", callback_data=f"follow_no_{export_id}")
            ],
            [
                InlineKeyboardButton(text="🤔 Нужна консультация", callback_data=f"follow_consult_{export_id}")
            ]
        ]
    )

def get_mailing_filters_keyboard():
    """Клавиатура фильтров для рассылки"""
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="👥 Все подписанные")],
            [KeyboardButton(text="📝 С анкетами")],
            [KeyboardButton(text="📭 Без анкет")],
            [KeyboardButton(text="🆕 За неделю")],
            [KeyboardButton(text="❌ Отмена")]
        ],
        resize_keyboard=True
    )

def get_mailing_feedback_keyboard(sent_message_id: int):
    """Клавиатура для обратной связи по рассылке"""
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="👍 Понравилось", callback_data=f"feedback_like_{sent_message_id}"),
                InlineKeyboardButton(text="👎 Не понравилось", callback_data=f"feedback_dislike_{sent_message_id}")
            ],
            [
                InlineKeyboardButton(text="💬 Комментарий", callback_data=f"feedback_comment_{sent_message_id}"),
                InlineKeyboardButton(text="🚫 Отписаться", callback_data=f"feedback_unsubscribe_{sent_message_id}")
            ]
        ]
    )

def get_subscription_management_keyboard(user_id: int, current_status: bool):
    """Клавиатура управления подпиской пользователя"""
    status_text = "✅ Подписан" if current_status else "❌ Отписан"
    
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"{status_text} - Переключить", 
                    callback_data=f"toggle_sub_{user_id}"
                )
            ],
            [
                InlineKeyboardButton(
                    text="📊 Статистика пользователя", 
                    callback_data=f"user_stats_{user_id}"
                )
            ]
        ]
    )

def get_manager_response_keyboard(message_id: int):
    """Клавиатура для ответа менеджеру"""
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="📞 Позвонить", callback_data=f"call_{message_id}")],
            [InlineKeyboardButton(text="💬 Написать в Telegram", callback_data=f"write_{message_id}")],
            [InlineKeyboardButton(text="✅ Обработано", callback_data=f"done_{message_id}")]
        ]
    )

def get_export_confirmation_keyboard(export_id: int):
    """Клавиатура подтверждения отправки выгрузки"""
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="✅ Подтвердить отправку", callback_data=f"confirm_export_{export_id}"),
                InlineKeyboardButton(text="❌ Отменить", callback_data=f"cancel_export_{export_id}")
            ]
        ]
    )

def get_export_user_input_keyboard():
    """Клавиатура для ввода ID пользователя"""
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="❌ Отмена")]
        ],
        resize_keyboard=True
    )

def get_export_notification_keyboard():
    """Клавиатура для уведомления о новой выгрузке"""
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="📊 Посмотреть мои выгрузки", callback_data="my_exports_callback")
            ]
        ]
    )

def get_request_contacts_keyboard(export_id: int):
    """Клавиатура для запроса контактов у пользователя"""
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="📝 Заполнить контакты для получения выгрузки", callback_data=f"fill_contacts_{export_id}")]
        ]
    )

# =========== СОСТОЯНИЯ ===========
class Questionnaire(StatesGroup):
    # Первая часть (1-4 вопросы)
    waiting_for_activity = State()
    waiting_for_region = State()
    waiting_for_budget = State()
    waiting_for_keywords = State()
    
    # Вторая часть (5-8 вопросы) - для получения выгрузки
    waiting_for_company = State()
    waiting_for_name = State()
    waiting_for_phone = State()
    waiting_for_email = State()

class ManagerDialog(StatesGroup):
    waiting_for_message = State()

class ManualMailing(StatesGroup):
    waiting_for_text = State()
    waiting_for_filter = State()
    waiting_for_confirmation = State()

class FeedbackComment(StatesGroup):
    waiting_for_comment = State()

class SendExport(StatesGroup):
    waiting_for_user_id = State()
    waiting_for_export_file = State()

class ExportContacts(StatesGroup):
    """Состояния для сбора контактов перед отправкой выгрузки"""
    waiting_for_company = State()
    waiting_for_name = State()
    waiting_for_phone = State()
    waiting_for_email = State()

# =========== ФУНКЦИЯ ОТПРАВКИ ЧАСТИЧНОЙ АНКЕТЫ АДМИНИСТРАТОРУ ===========
async def send_partial_questionnaire_to_admin(questionnaire_id: int, user_id: int, user_data: dict, username: str):
    """Отправка первой части анкеты администратору"""
    if not ADMIN_ID:
        logger.warning("ADMIN_ID не установлен, анкета не отправлена администратору")
        return
    
    try:
        admin_message = f"""
📋 <b>ЧАСТИЧНАЯ АНКЕТА #{questionnaire_id} (1-4 пункты)</b>

👤 <b>Пользователь:</b> @{username or 'без username'}
🆔 <b>Telegram ID:</b> {user_id}
📅 <b>Дата заполнения:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}

<b>Данные анкеты (1-4):</b>

<b>1. Сфера деятельности компании:</b>
{user_data.get('activity', 'Не указано')}

<b>2. Регионы работы:</b>
{user_data.get('region', 'Не указано')}

<b>3. Бюджет контрактов:</b>
{user_data.get('budget', 'Не указано')}

<b>4. Ключевые слова для поиска:</b>
{user_data.get('keywords', 'Не указано')}

<i>Пользователь ожидает выгрузку тендеров.
Для завершения анкеты нужны контакты (пункты 5-8).</i>
        """
        
        await bot.send_message(ADMIN_ID, admin_message, parse_mode=ParseMode.HTML)
        logger.info(f"Частичная анкета #{questionnaire_id} отправлена администратору")
        
    except Exception as e:
        logger.error(f"Ошибка при отправке частичной анкеты администратору: {e}")

# =========== ФУНКЦИЯ ОТПРАВКИ ПОЛНОЙ АНКЕТЫ АДМИНИСТРАТОРУ ===========
async def send_questionnaire_to_admin(questionnaire_id: int, user_id: int, user_data: dict, username: str, anketa_path: str = None):
    """Отправка заполненной анкеты администратору"""
    if not ADMIN_ID:
        logger.warning("ADMIN_ID не установлен, анкета не отправлена администратору")
        return
    
    try:
        admin_message = f"""
📋 <b>НОВАЯ АНКЕТА #{questionnaire_id}</b>

👤 <b>Пользователь:</b> @{username or 'без username'}
🆔 <b>Telegram ID:</b> {user_id}
📅 <b>Дата заполнения:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}

<b>Данные анкеты:</b>

<b>1. Сфера деятельности компании:</b>
{user_data.get('activity', 'Не указано')}

<b>2. Регионы работы:</b>
{user_data.get('region', 'Не указано')}

<b>3. Бюджет контрактов:</b>
{user_data.get('budget', 'Не указано')}

<b>4. Ключевые слова для поиска:</b>
{user_data.get('keywords', 'Не указано')}

<b>5. Название компании:</b>
{user_data.get('company_name', 'Не указано')}

<b>6. ФИО полностью:</b>
{user_data.get('full_name', 'Не указано')}

<b>7. Телефон для связи:</b>
{user_data.get('phone', 'Не указано')}

<b>8. Email для отправки тендеров:</b>
{user_data.get('email', 'Не указано')}

{'✅ <b>Заполнено в рабочее время</b>' if db.is_working_hours() else '⏰ <b>Заполнено в нерабочее время</b>'}
        """
        
        if anketa_path and os.path.exists(anketa_path):
            # Используем BufferedInputFile для отправки файла
            with open(anketa_path, 'rb') as f:
                file_data = f.read()
            
            input_file = BufferedInputFile(
                file_data,
                filename=f"Анкета_{questionnaire_id}_{username or 'user'}.docx"
            )
            
            await bot.send_document(
                ADMIN_ID,
                document=input_file,
                caption=admin_message,
                parse_mode=ParseMode.HTML
            )
            
            logger.info(f"Анкета #{questionnaire_id} с файлом отправлена администратору {ADMIN_ID}")
        else:
            await bot.send_message(ADMIN_ID, admin_message, parse_mode=ParseMode.HTML)
            logger.info(f"Анкета #{questionnaire_id} отправлена администратору {ADMIN_ID}")
        
    except Exception as e:
        logger.error(f"Ошибка при отправке анкеты администратору: {e}")

# =========== ФУНКЦИЯ ОТПРАВКИ ФАЙЛА ANKETA.DOCX ===========
async def send_anketa_file(message: types.Message, file_path: str):
    """Отправка файла анкеты пользователю"""
    try:
        # Проверяем наличие файла
        if not os.path.exists(file_path):
            logger.error(f"Файл не найден: {file_path}")
            await message.answer("❌ Файл анкеты не найден. Попробуйте позже.")
            return False
        
        # Используем BufferedInputFile для отправки файла
        with open(file_path, 'rb') as f:
            file_data = f.read()
        
        input_file = BufferedInputFile(
            file_data,
            filename="Анкета_Тритика_шаблон.docx"
        )
        
        await message.answer_document(
            document=input_file,
            caption=(
                "📄 <b>Шаблон анкеты для заполнения</b>\n\n"
                "Вы можете заполнить эту анкету и отправить нам:\n\n"
                "1. 📧 <b>На email:</b> info@tritika.ru\n"
                "2. 🤖 <b>Через бота:</b> кнопка 'Написать менеджеру'\n"
                "3. 👨‍💼 <b>Менеджеру в Telegram:</b> @tritikaru\n\n"
            ),
            parse_mode=ParseMode.HTML
        )
        
        logger.info(f"✅ Файл анкеты успешно отправлен пользователю {message.from_user.id}")
        return True
        
    except Exception as e:
        logger.error(f"Ошибка при отправке файла анкеты: {e}")
        
        # Если не удалось отправить файл, отправляем ссылку
        await message.answer(
            f"📄 <b>Шаблон анкеты для заполнения</b>\n\n"
            f"Скачать анкету можно по ссылке:\n{ANKETA_GITHUB_URL}\n\n"
            "Вы можете заполнить эту анкету и отправить нам:\n\n"
            "1. 📧 <b>На email:</b> info@tritika.ru\n"
            "2. 🤖 <b>Через бота:</b> кнопка 'Написать менеджеру'\n"
            "3. 👨‍💼 <b>Менеджеру в Telegram:</b> @tritikaru\n\n",
            parse_mode=ParseMode.HTML
        )
        return True  # Возвращаем True, так как пользователь получил ссылку

# =========== ФУНКЦИЯ ДЛЯ ОТПРАВКИ УВЕДОМЛЕНИЯ О НОВОЙ ВЫГРУЗКЕ ===========
async def send_export_notification_to_user(user_id: int, export_id: int, export_data: dict = None):
    """Отправка уведомления пользователю о новой выгрузке"""
    try:
        notification_message = f"""
📨 <b>НОВАЯ ВЫГРУЗКА ТЕНДЕРОВ #{export_id}</b>

📅 <b>Дата отправки:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}

<b>Выгрузка успешно подготовлена и отправлена!</b>

<i>Вы можете посмотреть все ваши выгрузки в разделе "📊 Мои выгрузки" или нажав кнопку ниже.</i>
"""
        
        if export_data:
            notification_message = f"""
📨 <b>НОВАЯ ВЫГРУЗКА ТЕНДЕРОВ #{export_id}</b>

🏢 <b>Компания:</b> {export_data.get('company_name', 'Ваша компания')}
🎯 <b>Сфера:</b> {export_data.get('activity', 'Не указано')}
📅 <b>Дата отправки:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}

<b>Выгрузка успешно подготовлена и отправлена!</b>

<i>Вы можете посмотреть все ваши выгрузки в разделе "📊 Мои выгрузки" или нажав кнопку ниже.</i>
"""
        
        await bot.send_message(
            user_id,
            notification_message,
            reply_markup=get_export_notification_keyboard(),
            parse_mode=ParseMode.HTML
        )
        
        logger.info(f"✅ Уведомление о выгрузке #{export_id} отправлено пользователю {user_id}")
        
    except Exception as e:
        logger.error(f"❌ Ошибка отправки уведомления о выгрузке пользователю {user_id}: {e}")

# =========== ФУНКЦИЯ ДЛЯ ОТПРАВКИ FOLLOW-UP СООБЩЕНИЙ ===========
async def send_follow_up_messages():
    """Отправка follow-up сообщений через 1 час после выгрузке"""
    try:
        exports = db.get_exports_for_followup()
        
        for export in exports:
            export_id = export['id']
            user_id = export['user_id']
            username = export['username'] or "Пользователь"
            
            try:
                await bot.send_message(
                    user_id,
                    f"📨 <b>Подборка тендеров отправлена!</b>\n\n"
                    f"Удалось ли найти что-то подходящее?",
                    reply_markup=get_follow_up_keyboard(export_id),
                    parse_mode=ParseMode.HTML
                )
                
                db.mark_followup_sent(export_id)
                
                logger.info(f"Follow-up отправлен пользователю {user_id} для выгрузки #{export_id}")
                
            except Exception as e:
                logger.error(f"Ошибка отправки follow-up пользователю {user_id}: {e}")
                
    except Exception as e:
        logger.error(f"Ошибка в send_follow_up_messages: {e}")

# =========== ФУНКЦИЯ ДЛЯ ПЛАНИРОВАНИЯ FOLLOW-UP ===========
async def schedule_follow_ups():
    """Планирование отправки follow-up сообщений"""
    while True:
        try:
            await send_follow_up_messages()
        except Exception as e:
            logger.error(f"Ошибка в schedule_follow_ups: {e}")
        
        await asyncio.sleep(300)

# =========== ФУНКЦИЯ ЗАПРОСА КОНТАКТОВ ДЛЯ ВЫГРУЗКИ ===========
async def send_contacts_request(user_id: int, export_id: int, export_data: dict):
    """Отправка запроса на контакты пользователю"""
    try:
        # Получаем частичную анкету пользователя для данных
        questionnaires = db.get_partial_questionnaires()
        user_questionnaire = None
        for q in questionnaires:
            if q['user_id'] == user_id:
                user_questionnaire = q
                break
        
        message_text = f"""
📋 <b>Мы проанализировали вашу анкету!</b>

✅ <b>Подготовили для Вас список тендеров</b>

🎯 <b>По вашим критериям:</b>
• Сфера: {user_questionnaire['activity'] if user_questionnaire else 'Ваша сфера деятельности'}
• Регионы: {user_questionnaire['region'] if user_questionnaire else 'Ваши регионы'}
• Бюджет: {user_questionnaire['budget'] if user_questionnaire else 'Ваш бюджет'}

📄 <b>Для получения выгрузки тендеров оставьте свои контакты:</b>

<i>Заполните оставшиеся пункты анкеты, и мы сразу отправим вам подготовленную выгрузку.</i>
"""
        
        # Создаем новую запись о запросе контактов
        db.create_contact_request(user_id, export_id)
        
        # Отправляем сообщение пользователю с кнопкой для заполнения контактов
        await bot.send_message(
            user_id,
            message_text,
            reply_markup=get_request_contacts_keyboard(export_id),
            parse_mode=ParseMode.HTML
        )
        
        logger.info(f"✅ Запрос контактов отправлен пользователю {user_id} для выгрузки #{export_id}")
        
        # Уведомляем администратора о запросе контактов
        if ADMIN_ID:
            try:
                user = db.get_user_by_id(user_id)
                user_name = f"{user['first_name']} {user['last_name'] or ''}" if user else f"ID: {user_id}"
                
                await bot.send_message(
                    ADMIN_ID,
                    f"📨 <b>Запрос контактов отправлен пользователю</b>\n\n"
                    f"👤 Пользователь: {user_name}\n"
                    f"🆔 ID: {user_id}\n"
                    f"📋 Выгрузка ID: {export_id}\n\n"
                    f"<i>Пользователю отправлен запрос на заполнение контактов для получения выгрузки.</i>",
                    parse_mode=ParseMode.HTML
                )
            except Exception as e:
                logger.error(f"Ошибка уведомления администратора: {e}")
        
        return True
        
    except Exception as e:
        logger.error(f"❌ Ошибка отправки запроса контактов пользователю {user_id}: {e}")
        return False

# =========== ФУНКЦИЯ ОТПРАВКИ ВЫГРУЗКИ ПОЛЬЗОВАТЕЛЮ ===========
async def send_export_to_user(export_id: int, export_data: dict):
    """Отправка выгрузки пользователю (если контакты уже есть)"""
    user_id = export_data['user_id']
    file_path = export_data['file_path']
    file_name = export_data['file_name']
    
    await send_export_file_to_user(user_id, file_path, file_name, export_id)

async def send_export_file_to_user(user_id: int, file_path: str, file_name: str, export_id: int):
    """Отправка файла выгрузки пользователю"""
    try:
        if file_path and os.path.exists(file_path):
            with open(file_path, 'rb') as f:
                file_data = f.read()
            
            input_file = BufferedInputFile(
                file_data,
                filename=file_name or "Выгрузка_тендеров.pdf"
            )
            
            await bot.send_document(
                user_id,
                document=input_file,
                caption=(
                    f"📨 <b>Ваша выгрузка тендеров #{export_id} готова!</b>\n\n"
                    f"📅 <b>Дата отправки:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n"
                    f"<i>Выгрузка успешно подготовлена по вашим критериям.</i>\n"
                    f"<i>Вы можете посмотреть ее в разделе '📊 Мои выгрузки'</i>"
                ),
                parse_mode=ParseMode.HTML
            )
            
            logger.info(f"✅ Файл выгрузки #{export_id} отправлен пользователю {user_id}")
            
            # Обновляем статус выгрузки
            db.mark_export_completed(export_id, "Автоматическая отправка")
        else:
            await bot.send_message(
                user_id,
                f"📨 <b>Ваша выгрузка тендеров #{export_id} готова!</b>\n\n"
                f"📅 <b>Дата отправки:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n"
                f"<i>Выгрузка была успешно подготовлена. "
                f"Вы можете посмотреть ее в разделе '📊 Мои выгрузки'.</i>",
                parse_mode=ParseMode.HTML
            )
            logger.info(f"✅ Уведомление о выгрузке отправлено пользователю {user_id} (без файла)")
            
    except Exception as e:
        logger.error(f"❌ Ошибка отправки выгрузки пользователю {user_id}: {e}")

# =========== ОБРАБОТЧИКИ КОМАНД ===========
@dp.message(Command("start"))
async def cmd_start(message: types.Message, state: FSMContext):
    """Обработчик команды /start"""
    await state.clear()
    
    user = message.from_user
    user_id = user.id
    
    db.add_user(user_id, user.username or "", user.first_name, user.last_name or "")
    
    is_admin = ADMIN_ID and user_id == ADMIN_ID
    
    if is_admin:
        await message.answer(
            "🛠️ <b>Панель администратора Тритика</b>\n\n"
            "Вы вошли как администратор бота.\n"
            "Используйте кнопки ниже для управления.",
            reply_markup=get_admin_keyboard(),
            parse_mode=ParseMode.HTML
        )
    else:
        await message.answer(
            "👋 <b>Привет! Я интеллектуальный ассистент компании Тритика.</b>\n\n"
            "Помогаю организацим находить выгодные тендеры. "
            "Хотите бесплатно получить подборку тендеров по вашей сфере? "
            "Вам надо лишь заполнить короткую анкету.",
            reply_markup=get_main_keyboard(),
            parse_mode=ParseMode.HTML
        )
    
    logger.info(f"Пользователь {user_id} нажал /start")

@dp.message(Command("help"))
async def cmd_help(message: types.Message):
    """Обработчик команды /help"""
    await message.answer(
        "🤖 <b>Помощь по боту Тритика:</b>\n\n"
        "<b>Основные команды:</b>\n"
        "/start - Главное меню\n"
        "/help - Эта справка\n"
        "/my_exports - Мои выгрузки\n\n"
        "<b>Основные функции:</b>\n"
        "• Заполнить анкету онлайн\n"
        "• Скачать анкету в Word\n"
        "• Написать менеджеру (отправить вопрос или заполненную анкету)\n"
        "• Получить бесплатную подборку тендеров\n"
        "• Консультация по участию в тендерах\n\n"
        "<b>Контакты поддержки:</b>\n"
        "📧 info@tritika.ru\n"
        "📱 +7 (904) 653-69-87",
        parse_mode=ParseMode.HTML
    )

@dp.message(Command("my_exports"))
async def cmd_my_exports(message: types.Message):
    """Мои выгрузки - показываем подробную информацию"""
    user_id = message.from_user.id
    
    exports = db.get_user_exports(user_id)
    
    if not exports:
        await message.answer(
            "📭 <b>У вас пока нет выгрузки тендеров.</b>\n\n"
            "Хотите получить бесплатную подборку? Заполните анкету!\n\n"
            "<i>После заполнения анкеты мы подготовим для вас подборку тендеров, и она появится в этом разделе.</i>",
            reply_markup=get_main_keyboard(),
            parse_mode=ParseMode.HTML
        )
        return
    
    response = f"📋 <b>Ваши выгрузки тендеров ({len(exports)}):</b>\n\n"
    
    for i, export in enumerate(exports, 1):
        date_str = export['sent_at'][:10] if export['sent_at'] else "??.??.????"
        status_icon = "✅" if export['status'] == 'completed' else "⏳" if export['status'] == 'pending' else "❌"
        status_text = {
            'completed': 'Отправлена',
            'pending': 'В обработке',
            'cancelled': 'Отменена'
        }.get(export['status'], export['status'])
        
        # Добавляем информацию о файле
        file_info = ""
        if export['file_name']:
            file_info = f"📄 {export['file_name']}"
        elif export['status'] == 'completed':
            file_info = "📝 Выгрузка отправлена (без файла)"
        
        response += f"<b>{i}. #{export['id']}</b>\n"
        response += f"   📅 <i>Дата запроса:</i> {date_str}\n"
        response += f"   📊 <i>Статус:</i> {status_icon} {status_text}\n"
        
        if file_info:
            response += f"   {file_info}\n"
        
        if export['follow_up_response']:
            response += f"   💬 <i>Ваш ответ:</i> {export['follow_up_response']}\n"
        
        response += "\n"
    
    response += "\n<i>Для обновления списка нажмите /my_exports или кнопку '📊 Мои выгрузки'</i>"
    
    await message.answer(response, parse_mode=ParseMode.HTML)

@dp.message(Command("admin"))
async def cmd_admin(message: types.Message, state: FSMContext):
    """Вход в админ-панель"""
    user_id = message.from_user.id
    
    if ADMIN_ID and user_id == ADMIN_ID:
        await state.clear()
        await message.answer(
            "🔐 <b>Вы авторизованы как администратор</b>",
            reply_markup=get_admin_keyboard(),
            parse_mode=ParseMode.HTML
        )
    else:
        await message.answer("⛔ У вас нет прав доступа к панели администратора.", parse_mode=ParseMode.HTML)

# =========== ОБРАБОТЧИК КОНТАКТА ИЗ ГЛАВНОГО МЕНЮ ===========
@dp.message(F.contact)
async def handle_main_phone_contact(message: types.Message):
    """Обработка контакта из главного меню"""
    user = message.from_user
    user_id = user.id
    
    phone = message.contact.phone_number
    
    # Сохраняем телефон в базе данных
    conn = sqlite3.connect("tenders.db")
    cursor = conn.cursor()
    
    cursor.execute('''
    UPDATE users 
    SET phone = ?
    WHERE user_id = ?
    ''', (phone, user_id))
    
    conn.commit()
    conn.close()
    
    await message.answer(
        f"✅ <b>Телефон сохранен!</b>\n\n"
        f"📱 <b>Ваш телефон:</b> {phone}\n\n"
        f"Теперь вы можете заполнить анкету для поиска тендеров.",
        reply_markup=get_main_keyboard(),
        parse_mode=ParseMode.HTML
    )

# =========== ОБРАБОТЧИКИ КНОПОК ===========
@dp.message(F.text == "📝 Заполнить анкету онлайн")
async def start_online_questionnaire(message: types.Message, state: FSMContext):
    """Начало заполнения анкеты онлайн - БЕЗ показа порядка"""
    await state.clear()
    
    await message.answer(
        "📝 <b>Заполнение анкеты онлайн</b>\n\n"
        "Введите <b>сферу деятельности вашей компании</b>:\n"
        "<i>Пример: строительство, IT-услуги, поставка продуктов питания</i>",
        reply_markup=get_cancel_keyboard(),
        parse_mode=ParseMode.HTML
    )
    await state.set_state(Questionnaire.waiting_for_activity)

@dp.message(F.text == "📥 Скачать анкету в Word")
async def download_questionnaire(message: types.Message, state: FSMContext):
    """Скачать анкету в Word"""
    await state.clear()
    
    await message.answer("📄 <b>Отправляю вам шаблон анкеты...</b>", parse_mode=ParseMode.HTML)
    
    # Проверяем наличие файла локально
    if not os.path.exists(ANKETA_LOCAL_PATH):
        logger.warning(f"Файл анкеты не найден локально: {ANKETA_LOCAL_PATH}")
        # Пробуем скачать с GitHub
        success = await download_anketa_file()
        if not success:
            # Если не удалось скачать, отправляем ссылку на GitHub
            await message.answer(
                f"📄 <b>Шаблон анкеты для заполнения</b>\n\n"
                f"Скачать анкету можно по ссылке:\n{ANKETA_GITHUB_URL}\n\n"
                "Вы можете заполнить эту анкету и отправить нам:\n\n"
                "1. 📧 <b>На email:</b> info@tritika.ru\n"
                "2. 🤖 <b>Через бота:</b> кнопка '📤 Написать менеджеру'\n"
                "3. 👨‍💼 <b>Менеджеру в Telegram:</b> @tritikaru\n\n",
                reply_markup=get_main_keyboard(),
                parse_mode=ParseMode.HTML
            )
            return
    
    # Отправляем файл анкеты
    sent = await send_anketa_file(message, ANKETA_LOCAL_PATH)
    
    if sent:
        await message.answer(
            "📝 <b>Что дальше?</b>\n\n"
            "Или вы можете заполнить анкету прямо здесь через <b>'Заполнить анкету онлайн'</b>",
            reply_markup=get_main_keyboard(),
            parse_mode=ParseMode.HTML
        )
    else:
        await message.answer(
            "❌ Не удалось отправить файл анкеты. Попробуйте позже или свяжитесь с поддержкой.",
            reply_markup=get_main_keyboard(),
            parse_mode=ParseMode.HTML
        )

@dp.message(F.text == "📤 Написать менеджеру")
async def start_manager_dialog(message: types.Message, state: FSMContext):
    """Начало диалога с менеджеру"""
    await state.set_state(ManagerDialog.waiting_for_message)
    await message.answer(
        "💬 <b>Напишите ваше сообщение менеджеру</b>\n\n"
        "Вы можете отправить:\n"
        "• Текст с вопросом\n"
        "• Заполненную анкету (файл Word)\n"
        "• Документы\n"
        "• Фотографии\n\n"
        "<i>Мы получим ваше сообщение и ответим в ближайшее время.</i>",
        reply_markup=get_cancel_keyboard(),
        parse_mode=ParseMode.HTML
    )

@dp.message(F.text == "📊 Мои выгрузки")
async def my_exports_button(message: types.Message):
    """Мои выгрузки через кнопку"""
    await cmd_my_exports(message)

@dp.message(F.text == "📞 Контакты")
async def show_contacts(message: types.Message):
    """Показать контакты"""
    await message.answer(
        "📞 <b>Контакты компании Тритика</b>\n\n"
        "<b>Техническая поддержка:</b>\n"
        "• Телефон: +7 (904) 653-69-87\n"
        "• Email: info@tritika.ru\n"
        "• Telegram: @tritikaru\n\n"
        "<b>Время работы:</b>\n"
        "Пн-Чт: 8:30-17:30\n"
        "Пт: 8:30-16:30\n"
        "Сб-Вс: выходные",
        parse_mode=ParseMode.HTML
    )

@dp.message(F.text == "ℹ️ Помощь")
async def show_help(message: types.Message):
    """Показать помощь"""
    await cmd_help(message)

@dp.message(F.text == "❌ Отмена")
async def cancel_action(message: types.Message, state: FSMContext):
    """Отмена действия"""
    current_state = await state.get_state()
    
    if current_state in [ManagerDialog.waiting_for_message, 
                         ManualMailing.waiting_for_text,
                         ManualMailing.waiting_for_filter,
                         ManualMailing.waiting_for_confirmation,
                         FeedbackComment.waiting_for_comment,
                         SendExport.waiting_for_user_id,
                         SendExport.waiting_for_export_file,
                         ExportContacts.waiting_for_company,
                         ExportContacts.waiting_for_name,
                         ExportContacts.waiting_for_phone,
                         ExportContacts.waiting_for_email]:
        await state.clear()
        is_admin = ADMIN_ID and message.from_user.id == ADMIN_ID
        
        if is_admin:
            await message.answer("❌ Действие отменено", reply_markup=get_admin_keyboard(), parse_mode=ParseMode.HTML)
        else:
            await message.answer(
                "❌ Действие отменено.\n\n"
                "Вы можете выбрать другое действие.",
                reply_markup=get_main_keyboard(),
                parse_mode=ParseMode.HTML
            )
    elif current_state in Questionnaire.__states__:
        await state.clear()
        await message.answer(
            "❌ Заполнение анкеты отменено.\n\n"
            "Вы можете начать заполнение заново в любое время.",
            reply_markup=get_main_keyboard(),
            parse_mode=ParseMode.HTML
        )
    else:
        await state.clear()
        await message.answer(
            "❌ Действие отменено.",
            reply_markup=get_main_keyboard(),
            parse_mode=ParseMode.HTML
        )

# =========== ДИАЛОГ С МЕНЕДЖЕРОМ ===========
@dp.message(ManagerDialog.waiting_for_message)
async def process_manager_message(message: types.Message, state: FSMContext):
    """Обработка сообщения для менеджера"""
    user = message.from_user
    user_id = user.id
    
    message_type = "text"
    file_id = None
    file_name = None
    
    if message.document:
        message_type = "document"
        file_id = message.document.file_id
        file_name = message.document.file_name
        message_text = f"Документ: {message.document.file_name}"
    elif message.photo:
        message_type = "photo"
        file_id = message.photo[-1].file_id
        message_text = "Фотография"
    elif message.text:
        message_text = message.text
    else:
        await message.answer("❌ Извините, я могу принимать только текст, документы и фотографии.", parse_mode=ParseMode.HTML)
        return
    
    message_id = db.save_manager_message(user_id, message_type, message_text, file_id, file_name)
    
    if ADMIN_ID:
        try:
            admin_message = f"📩 <b>НОВОЕ СООБЩЕНИЕ ОТ ПОЛЬЗОВАТЕЛЯ</b>\n\n"
            admin_message += f"👤 <b>Пользователь:</b> @{user.username or 'без username'}\n"
            admin_message += f"🆔 <b>ID:</b> {user_id}\n"
            admin_message += f"👤 <b>Имя:</b> {user.first_name} {user.last_name or ''}\n"
            admin_message += f"📅 <b>Время:</b> {datetime.now().strftime('%H:%M %d.%m.%Y')}\n"
            admin_message += f"📝 <b>Тип:</b> {message_type}\n\n"
            
            if message_type == "text":
                admin_message += f"💬 <b>Сообщение:</b>\n{message_text[:500]}"
                if len(message_text) > 500:
                    admin_message += "..."
            
            elif message_type == "document":
                admin_message += f"📎 <b>Документ:</b> {file_name}\n"
                admin_message += f"💬 <b>Сообщение:</b>\n{message_text}"
                
            elif message_type == "photo":
                admin_message += f"🖼 <b>Фотография</b>\n"
                admin_message += f"💬 <b>Сообщение:</b>\n{message_text}"
            
            keyboard = get_manager_response_keyboard(message_id)
            await bot.send_message(ADMIN_ID, admin_message, reply_markup=keyboard, parse_mode=ParseMode.HTML)
            
            if file_id:
                if message_type == "document":
                    await bot.send_document(ADMIN_ID, file_id, caption=f"Документ от пользователя {user_id}", parse_mode=ParseMode.HTML)
                elif message_type == "photo":
                    await bot.send_photo(ADMIN_ID, file_id, caption=f"Фото от пользователя {user_id}", parse_mode=ParseMode.HTML)
            
        except Exception as e:
            logger.error(f"Не удалось отправить уведомление админу: {e}")
    
    await message.answer(
        "✅ <b>Ваше сообщение отправлено менеджеру!</b>\n\n"
        "Мы получили ваше сообщение и свяжемся с вами в ближайшее время.\n\n"
        "<i>Обычно мы отвечаем в течение 15 минут в рабочее время.</i>",
        reply_markup=get_main_keyboard(),
        parse_mode=ParseMode.HTML
    )
    
    await state.clear()

# =========== CALLBACK ОБРАБОТЧИКИ ДЛЯ АДМИНА ===========
@dp.callback_query(F.data.startswith("call_"))
async def handle_call_callback(callback: types.CallbackQuery):
    """Обработка кнопки "Позвонить" для сообщения менеджеру"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    message_id = int(callback.data.split("_")[1])
    
    conn = sqlite3.connect("tenders.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('''
    SELECT mm.*, u.phone, u.first_name, u.last_name 
    FROM manager_messages mm
    JOIN users u ON mm.user_id = u.user_id
    WHERE mm.id = ?
    ''', (message_id,))
    
    message = cursor.fetchone()
    conn.close()
    
    if not message:
        await callback.answer("Сообщение не найдено", show_alert=True)
        return
    
    phone = message['phone']
    user_name = f"{message['first_name']} {message['last_name'] or ''}".strip()
    
    if phone:
        response = f"📞 <b>Телефон пользователя:</b> {phone}\n"
        response += f"👤 <b>Имя:</b> {user_name}\n"
        response += f"🆔 <b>ID:</b> {message['user_id']}\n"
        response += f"📅 <b>Время сообщения:</b> {message['created_at'][:19]}"
    else:
        response = "❌ У пользователя не указан телефон в анкете."
    
    await callback.message.answer(response, parse_mode=ParseMode.HTML)
    await callback.answer()

@dp.callback_query(F.data.startswith("write_"))
async def handle_write_callback(callback: types.CallbackQuery):
    """Обработка кнопки "Написать в Telegram" для сообщения менеджеру"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    message_id = int(callback.data.split("_")[1])
    
    conn = sqlite3.connect("tenders.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('''
    SELECT mm.*, u.username, u.first_name, u.last_name 
    FROM manager_messages mm
    JOIN users u ON mm.user_id = u.user_id
    WHERE mm.id = ?
    ''', (message_id,))
    
    message = cursor.fetchone()
    conn.close()
    
    if not message:
        await callback.answer("Сообщение не найдено", show_alert=True)
        return
    
    username = message['username']
    user_name = f"{message['first_name']} {message['last_name'] or ''}".strip()
    
    if username:
        response = f"✏️ <b>Написать пользователю:</b>\n"
        response += f"👤 <b>Username:</b> @{username}\n"
        response += f"👤 <b>Имя:</b> {user_name}\n"
        response += f"🆔 <b>ID:</b> {message['user_id']}\n"
        response += f"🔗 <b>Ссылка:</b> https://t.me/{username}"
    else:
        response = f"✏️ <b>Написать пользователю:</b>\n"
        response += f"👤 <b>Имя:</b> {user_name}\n"
        response += f"🆔 <b>ID:</b> {message['user_id']}\n"
        response += f"🔗 <b>Ссылка:</b> tg://user?id={message['user_id']}"
    
    await callback.message.answer(response, parse_mode=ParseMode.HTML)
    await callback.answer()

@dp.callback_query(F.data.startswith("done_"))
async def handle_done_callback(callback: types.CallbackQuery):
    """Обработка кнопки "Обработано" для сообщения менеджеру"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    message_id = int(callback.data.split("_")[1])
    
    conn = sqlite3.connect("tenders.db")
    cursor = conn.cursor()
    
    cursor.execute('''
    UPDATE manager_messages 
    SET processed = 1
    WHERE id = ?
    ''', (message_id,))
    
    conn.commit()
    conn.close()
    
    await callback.message.edit_text(
        callback.message.text + "\n\n✅ <b>ОБРАБОТАНО</b>",
        reply_markup=None,
        parse_mode=ParseMode.HTML
    )
    
    await callback.answer("Сообщение отмечено как обработанное")

# =========== CALLBACK ДЛЯ ПОЛЬЗОВАТЕЛЯ ===========
@dp.callback_query(F.data == "my_exports_callback")
async def handle_my_exports_callback(callback: types.CallbackQuery):
    """Обработка кнопки "Посмотреть мои выгрузки" """
    await cmd_my_exports(callback.message)
    await callback.answer()

# =========== АДМИН ПАНЕЛЬ ===========
@dp.message(F.text == "📊 Частичные анкеты")
async def show_partial_questionnaires(message: types.Message):
    """Показать частичные анкеты (только 1-4 пункты)"""
    if not ADMIN_ID or message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Доступ запрещен", parse_mode=ParseMode.HTML)
        return
    
    questionnaires = db.get_partial_questionnaires()
    
    if not questionnaires:
        await message.answer("📭 Частичных анкет нет", parse_mode=ParseMode.HTML)
        return
    
    response = f"📋 <b>Частичные анкеты (ожидают выгрузку) ({len(questionnaires)}):</b>\n\n"
    
    for i, q in enumerate(questionnaires, 1):
        date_str = q['created_at'][:16] if q['created_at'] else "??.?? ??:??"
        response += f"<b>{i}. #{q['id']}</b>\n"
        response += f"👤 @{q['username'] or 'без username'}\n"
        response += f"🎯 {q['activity'][:30]}...\n"
        response += f"📍 {q['region'][:30]}...\n"
        response += f"💰 {q['budget'][:30]}...\n"
        response += f"⏰ {date_str}\n\n"
    
    await message.answer(response, parse_mode=ParseMode.HTML)

@dp.message(F.text == "📤 Отправить выгрузку")
async def start_send_export(message: types.Message, state: FSMContext):
    """Начало отправки выгрузки пользователю"""
    if not ADMIN_ID or message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Доступ запрещен", parse_mode=ParseMode.HTML)
        return
    
    await state.set_state(SendExport.waiting_for_user_id)
    await message.answer(
        "📤 <b>Отправка выгрузки пользователю</b>\n\n"
        "Введите Telegram ID пользователя (число):\n"
        "<i>ID можно получить из списка пользователей или из сообщений</i>",
        reply_markup=get_export_user_input_keyboard(),
        parse_mode=ParseMode.HTML
    )

@dp.message(SendExport.waiting_for_user_id)
async def process_export_user_id(message: types.Message, state: FSMContext):
    """Обработка ID пользователя для отправки выгрузки"""
    if message.text == "❌ Отмена":
        await state.clear()
        await message.answer("❌ Отправка выгрузки отменена", reply_markup=get_admin_keyboard(), parse_mode=ParseMode.HTML)
        return
    
    if not message.text.isdigit():
        await message.answer("❌ Пожалуйста, введите числовой Telegram ID пользователя", parse_mode=ParseMode.HTML)
        return
    
    user_id = int(message.text)
    
    # Проверяем существование пользователя в базе
    user = db.get_user_by_id(user_id)
    
    if not user:
        # Если пользователя нет в базе, добавляем его
        try:
            # Пробуем получить информацию о пользователе через Telegram API
            chat = await bot.get_chat(user_id)
            db.add_user(user_id, chat.username or "", chat.first_name or "", chat.last_name or "")
            user = db.get_user_by_id(user_id)
            
            if not user:
                await message.answer(f"❌ Не удалось добавить пользователя {user_id} в базу данных", parse_mode=ParseMode.HTML)
                return
        except Exception as e:
            # Если не удалось получить информацию, все равно добавляем с минимальными данными
            db.add_user(user_id, "", f"Пользователь_{user_id}", "")
            user = db.get_user_by_id(user_id)
    
    await state.update_data(user_id=user_id)
    await state.set_state(SendExport.waiting_for_export_file)
    
    user_name = f"{user['first_name']} {user['last_name'] or ''}".strip()
    username = f"@{user['username']}" if user['username'] else "без username"
    
    await message.answer(
        f"✅ <b>Пользователь найден/добавлен</b>\n\n"
        f"👤 <b>Пользователь:</b> {user_name}\n"
        f"📱 <b>Username:</b> {username}\n"
        f"🆔 <b>Telegram ID:</b> {user_id}\n\n"
        f"📤 <b>Отправьте файл с выгрузкой тендеров:</b>\n\n"
        f"<i>Поддерживаются файлы: PDF, Excel, Word, ZIP, RAR, TXT</i>\n"
        f"<i>Или отправьте текст для создания выгрузки без файла</i>",
        reply_markup=get_cancel_keyboard(),
        parse_mode=ParseMode.HTML
    )

@dp.message(SendExport.waiting_for_export_file)
async def process_export_file(message: types.Message, state: FSMContext):
    """Обработка файла выгрузки"""
    if message.text == "❌ Отмена":
        await state.clear()
        await message.answer("❌ Отправка выгрузки отменена", reply_markup=get_admin_keyboard(), parse_mode=ParseMode.HTML)
        return
    
    data = await state.get_data()
    user_id = data.get('user_id')
    
    if not user_id:
        await message.answer("❌ Ошибка: ID пользователя не найден", parse_mode=ParseMode.HTML)
        await state.clear()
        return
    
    user = db.get_user_by_id(user_id)
    if not user:
        await message.answer("❌ Пользователь не найден", parse_mode=ParseMode.HTML)
        await state.clear()
        return
    
    file_id = None
    file_name = None
    file_path = None
    text_export = None
    
    # Если пользователь отправил документ
    if message.document:
        file_id = message.document.file_id
        file_name = message.document.file_name
        
        try:
            file = await bot.get_file(file_id)
            file_path = file.file_path
            
            # Создаем уникальное имя файла для экспорта
            timestamp = int(datetime.now().timestamp())
            safe_file_name = "".join([c if c.isalnum() or c in "._-" else "_" for c in file_name])
            export_filename = f"export_{user_id}_{timestamp}_{safe_file_name}"
            export_path = os.path.join(EXPORTS_DIR, export_filename)
            
            # Скачиваем файл в папку exports
            await bot.download_file(file_path, export_path)
            
            # Создаем запись о выгрузке
            export_id = db.create_tender_export(
                user_id,
                export_path,
                file_name
            )
            
        except Exception as e:
            logger.error(f"❌ Ошибка обработки файла выгрузки: {e}")
            await message.answer(f"❌ Ошибка обработки файла: {e}", parse_mode=ParseMode.HTML)
            await state.clear()
            return
    
    # Если пользователь отправил текст
    elif message.text:
        text_export = message.text
        
        # Создаем текстовый файл с выгрузкой
        try:
            timestamp = int(datetime.now().timestamp())
            export_filename = f"export_{user_id}_{timestamp}_text.txt"
            export_path = os.path.join(EXPORTS_DIR, export_filename)
            
            with open(export_path, 'w', encoding='utf-8') as f:
                f.write(text_export)
            
            # Создаем запись о выгрузке
            export_id = db.create_tender_export(
                user_id,
                export_path,
                "Выгрузка_тендеров.txt"
            )
            
        except Exception as e:
            logger.error(f"❌ Ошибка создания текстовой выгрузки: {e}")
            # Создаем выгрузку без файла
            export_id = db.create_tender_export_without_file(user_id)
    
    else:
        await message.answer("❌ Пожалуйста, отправьте файл или текст с выгрузкой", parse_mode=ParseMode.HTML)
        return
    
    keyboard = get_export_confirmation_keyboard(export_id)
    
    user_name = f"{user['first_name']} {user['last_name'] or ''}".strip()
    username = f"@{user['username']}" if user['username'] else "без username"
    
    if file_name:
        file_info = f"📄 <b>Файл:</b> {file_name}"
    elif text_export:
        file_info = f"📝 <b>Текстовая выгрузка:</b> {len(text_export)} символов"
    else:
        file_info = "📝 <b>Выгрузка без файла</b>"
    
    await message.answer(
        f"📤 <b>Подтверждение отправки выгрузки</b>\n\n"
        f"{file_info}\n"
        f"👤 <b>Пользователь:</b> {user_name}\n"
        f"📱 <b>Username:</b> {username}\n"
        f"🆔 <b>Telegram ID:</b> {user_id}\n"
        f"🆔 <b>ID выгрузки:</b> {export_id}\n\n"
        f"<i>Подтвердите отправку выгрузки пользователю.</i>",
        reply_markup=keyboard,
        parse_mode=ParseMode.HTML
    )
    
    await state.clear()

@dp.callback_query(F.data.startswith("confirm_export_"))
async def handle_confirm_export(callback: types.CallbackQuery):
    """Подтверждение отправки выгрузки - с проверкой контактов"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    export_id = int(callback.data.split("_")[2])
    
    export = db.get_export_by_id(export_id)
    
    if not export:
        await callback.answer("Выгрузка не найдена", show_alert=True)
        return
    
    try:
        user_id = export['user_id']
        
        # Проверяем, есть ли у пользователя полная анкета с контактами
        if db.has_complete_questionnaire(user_id):
            # Если контакты есть, отправляем выгрузку сразу
            await send_export_to_user(export_id, export)
            
            await callback.message.edit_text(
                callback.message.text + "\n\n✅ <b>ВЫГРУЗКА ОТПРАВЛЕНА ПОЛЬЗОВАТЕЛЮ</b>",
                reply_markup=None,
                parse_mode=ParseMode.HTML
            )
            
            await callback.message.answer(
                f"✅ <b>Выгрузка #{export_id} успешно отправлена пользователю</b>\n\n"
                f"👤 Пользователь: {export['first_name']} {export['last_name'] or ''}\n"
                f"📱 Username: @{export['username'] or 'без username'}\n"
                f"🆔 Telegram ID: {user_id}\n"
                f"{'📄 Файл: ' + export['file_name'] if export['file_name'] else '📝 Без файла'}\n\n"
                f"<i>Пользователь получил уведомление о новой выгрузке.</i>",
                parse_mode=ParseMode.HTML
            )
            
            logger.info(f"✅ Выгрузка #{export_id} отправлена пользователю {user_id}")
        else:
            # Если контактов нет, отправляем запрос на контакты
            success = await send_contacts_request(user_id, export_id, export)
            
            if success:
                await callback.message.edit_text(
                    callback.message.text + "\n\n📨 <b>ЗАПРОС КОНТАКТОВ ОТПРАВЛЕН</b>",
                    reply_markup=None,
                    parse_mode=ParseMode.HTML
                )
                
                await callback.message.answer(
                    f"📨 <b>Запрос контактов отправлен пользователю</b>\n\n"
                    f"👤 Пользователь ID: {user_id}\n"
                    f"📱 Username: @{export['username'] or 'без username'}\n"
                    f"🆔 Выгрузка ID: {export_id}\n\n"
                    f"<i>Пользователь получил сообщение с предложением заполнить контакты для получения выгрузки.</i>",
                    parse_mode=ParseMode.HTML
                )
            else:
                await callback.message.edit_text(
                    callback.message.text + "\n\n❌ <b>ОШИБКА ОТПРАВКИ ЗАПРОСА КОНТАКТОВ</b>",
                    reply_markup=None,
                    parse_mode=ParseMode.HTML
                )
                
                await callback.message.answer(
                    f"❌ <b>Не удалось отправить запрос контактов пользователю</b>\n\n"
                    f"👤 Пользователь ID: {user_id}\n"
                    f"📱 Username: @{export['username'] or 'без username'}\n"
                    f"🆔 Выгрузка ID: {export_id}\n\n"
                    f"<i>Возможно, пользователь заблокировал бота или произошла ошибка отправки.</i>",
                    parse_mode=ParseMode.HTML
                )
        
        await callback.answer()
        
    except Exception as e:
        logger.error(f"❌ Ошибка в handle_confirm_export: {e}")
        await callback.answer(f"❌ Ошибка: {str(e)[:100]}", show_alert=True)

@dp.callback_query(F.data.startswith("cancel_export_"))
async def handle_cancel_export(callback: types.CallbackQuery):
    """Отмена отправки выгрузки"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    export_id = int(callback.data.split("_")[2])
    
    export = db.get_export_by_id(export_id)
    if export and export['file_path'] and os.path.exists(export['file_path']):
        try:
            os.remove(export['file_path'])
            logger.info(f"✅ Файл выгрузки удален: {export['file_path']}")
        except Exception as e:
            logger.error(f"Не удалось удалить файл при отмене выгрузки: {e}")
    
    conn = sqlite3.connect("tenders.db")
    cursor = conn.cursor()
    cursor.execute('UPDATE tender_exports SET status = "cancelled" WHERE id = ?', (export_id,))
    conn.commit()
    conn.close()
    
    await callback.message.edit_text(
        callback.message.text + "\n\n❌ <b>ОТПРАВКА ОТМЕНЕНА</b>",
        reply_markup=None,
        parse_mode=ParseMode.HTML
    )
    
    await callback.answer("Отправка выгрузки отменена")

@dp.message(F.text == "📈 Статистика")
async def show_statistics(message: types.Message):
    """Показать статистику"""
    if not ADMIN_ID or message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Доступ запрещен", parse_mode=ParseMode.HTML)
        return
    
    stats = db.get_statistics(14)
    
    # Получаем количество частичных и полных анкет
    partial = len(db.get_partial_questionnaires())
    complete = len(db.get_complete_questionnaires())
    
    response = f"""
📊 <b>Статистика за 2 недели</b>

👥 <b>Пользователи:</b>
• Новых пользователей: {stats['new_users']}
• Частичных анкет: {partial}
• Полных анкет: {complete}
• С подпиской: {stats['subscribed_users']}
• Без подписки: {stats['unsubscribed_users']}

📋 <b>Выгрузки:</b>
• Выполненных выгрузок: {stats['exports_completed']}

💬 <b>Сообщения менеджеру:</b>
• Всего сообщений: {stats['manager_messages']}

📨 <b>Ручные рассылки:</b>
• Количество рассылок: {stats['mailings_count']}
• Отправлено сообщений: {stats['mailings_sent']}
• Получено отзывов: {stats['mailings_feedback']}

📅 <b>Дата отчета:</b>
{datetime.now().strftime('%d.%m.%Y %H:%M')}
"""
    
    await message.answer(response, parse_mode=ParseMode.HTML)

# =========== ОБРАБОТЧИКИ FOLLOW-UP СООБЩЕНИЙ ===========
@dp.callback_query(F.data.startswith("follow_"))
async def handle_follow_up_response(callback: types.CallbackQuery):
    """Обработка ответа на follow-up сообщение"""
    try:
        parts = callback.data.split("_")
        response_type = parts[1]
        export_id = int(parts[2])
        
        user_id = callback.from_user.id
        username = callback.from_user.username or "без username"
        
        response_map = {
            "yes": "Да, нашел подходящее",
            "no": "Нет, не нашел",
            "consult": "Нужна консультация"
        }
        
        response_text = response_map.get(response_type, "Неизвестно")
        db.save_followup_response(export_id, response_text)
        
        thank_you_text = {
            "yes": "Отлично! Мы рады, что вы нашли подходящие тендеры. 🎉",
            "no": "Жаль, что не нашли подходящее. Мы можем сделать более точную подборку. 📊",
            "consult": "Хорошо! Наш менеджер свяжется с вами для консультации. 👨‍💼"
        }
        
        await callback.message.edit_text(
            callback.message.text + f"\n\n✅ <b>Спасибо за ваш ответ!</b>\n{thank_you_text.get(response_type, '')}",
            reply_markup=None,
            parse_mode=ParseMode.HTML
        )
        
        if ADMIN_ID:
            try:
                export = db.get_export_by_id(export_id)
                if export:
                    await bot.send_message(
                        ADMIN_ID,
                        f"📨 <b>ПОЛЬЗОВАТЕЛЬ ОТВЕТИЛ НА FOLLOW-UP</b>\n\n"
                        f"👤 Пользователь: @{username}\n"
                        f"🆔 ID: {user_id}\n"
                        f"💬 Ответ: {response_text}\n"
                        f"📅 Время: {datetime.now().strftime('%H:%M %d.%m.%Y')}",
                        parse_mode=ParseMode.HTML
                    )
            except Exception as e:
                logger.error(f"Не удалось уведомить админа о follow-up: {e}")
        
        await callback.answer()
        
    except Exception as e:
        logger.error(f"Ошибка обработки follow-up ответа: {e}")
        await callback.answer("Произошла ошибка", show_alert=True)

# =========== УПРАВЛЕНИЕ ПОДПИСКАМИ ===========
@dp.message(F.text == "👥 Управление подписками")
async def manage_subscriptions(message: types.Message):
    """Управление подписками пользователей"""
    if not ADMIN_ID or message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Доступ запрещен", parse_mode=ParseMode.HTML)
        return
    
    users = db.get_all_users_with_subscription(30)
    
    if not users:
        await message.answer("👥 Пользователей нет", parse_mode=ParseMode.HTML)
        return
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[])
    
    for user in users:
        status_icon = "✅" if user['mailing_subscribed'] else "❌"
        has_anketa = "📋" if user['has_filled_questionnaire'] else "📭"
        
        button_text = f"{status_icon} {has_anketa} {user['first_name']}"
        if user['last_name']:
            button_text += f" {user['last_name']}"
        
        if user['username']:
            button_text += f" (@{user['username']})"
        
        keyboard.inline_keyboard.append([
            InlineKeyboardButton(
                text=button_text[:50],
                callback_data=f"manage_user_{user['user_id']}"
            )
        ])
    
    keyboard.inline_keyboard.append([
        InlineKeyboardButton(text="✅ Только подписанные", callback_data="filter_subscribed"),
        InlineKeyboardButton(text="❌ Только отписанные", callback_data="filter_unsubscribed")
    ])
    
    keyboard.inline_keyboard.append([
        InlineKeyboardButton(text="📊 Статистика подписок", callback_data="subscription_stats"),
        InlineKeyboardButton(text="🔄 Обновить список", callback_data="refresh_subs")
    ])
    
    await message.answer(
        "👥 <b>Управление подписками на рассылку</b>\n\n"
        "Выберите пользователя для управления его подпиской:\n\n"
        "<b>Легенда:</b>\n"
        "✅ - подписан на рассылку\n"
        "❌ - отписан от рассылки\n"
        "📋 - заполнил анкету\n"
        "📭 - без анкеты",
        reply_markup=keyboard,
        parse_mode=ParseMode.HTML
    )

@dp.callback_query(F.data.startswith("manage_user_"))
async def handle_manage_user(callback: types.CallbackQuery):
    """Обработка выбора пользователя для управления подпиской"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    user_id = int(callback.data.split("_")[2])
    
    user_info = db.get_user_mailing_status(user_id)
    
    if not user_info:
        await callback.answer("Пользователь не найден", show_alert=True)
        return
    
    keyboard = get_subscription_management_keyboard(user_id, user_info['subscribed'])
    
    user_name = f"{user_info['first_name']} {user_info['last_name'] or ''}".strip()
    username = f"@{user_info['username']}" if user_info['username'] else "без username"
    
    await callback.message.edit_text(
        f"👤 <b>Управление подписки пользователя</b>\n\n"
        f"<b>Пользователь:</b> {user_name}\n"
        f"<b>Username:</b> {username}\n"
        f"<b>ID:</b> {user_id}\n"
        f"<b>Текущий статус:</b> {'✅ Подписан на рассылку' if user_info['subscribed'] else '❌ Отписан от рассылки'}\n\n"
        f"<i>Используйте кнопки ниже для управления:</i>",
        reply_markup=keyboard,
        parse_mode=ParseMode.HTML
    )
    
    await callback.answer()

@dp.callback_query(F.data.startswith("toggle_sub_"))
async def handle_toggle_subscription(callback: types.CallbackQuery):
    """Переключение статуса подписки"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    user_id = int(callback.data.split("_")[2])
    
    new_status = db.toggle_user_mailing_subscription(user_id)
    
    if new_status is None:
        await callback.answer("Ошибка при изменении подписки", show_alert=True)
        return
    
    user_info = db.get_user_mailing_status(user_id)
    
    if not user_info:
        await callback.answer("Пользователь не найден", show_alert=True)
        return
    
    keyboard = get_subscription_management_keyboard(user_id, new_status)
    
    user_name = f"{user_info['first_name']} {user_info['last_name'] or ''}".strip()
    
    await callback.message.edit_text(
        f"👤 <b>Управление подписки пользователя</b>\n\n"
        f"<b>Пользователь:</b> {user_name}\n"
        f"<b>ID:</b> {user_id}\n"
        f"<b>Текущий статус:</b> {'✅ Подписан на рассылку' if new_status else '❌ Отписан от рассылки'}\n\n"
        f"<i>Статус успешно обновлен!</i>",
        reply_markup=keyboard,
        parse_mode=ParseMode.HTML
    )
    
    await callback.answer(f"Статус подписки изменен: {'✅ Подписан' if new_status else '❌ Отписан'}")

@dp.callback_query(F.data.startswith("user_stats_"))
async def handle_user_stats(callback: types.CallbackQuery):
    """Показать статистику пользователя"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    user_id = int(callback.data.split("_")[2])
    
    conn = sqlite3.connect("tenders.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('''
    SELECT u.*, 
           COUNT(DISTINCT q.id) as questionnaire_count,
           COUNT(DISTINCT te.id) as export_count,
           COUNT(DISTINCT mm.id) as message_count,
           COUNT(DISTINCT mf.id) as feedback_count
    FROM users u
    LEFT JOIN questionnaires q ON u.user_id = q.user_id
    LEFT JOIN tender_exports te ON u.user_id = te.user_id
    LEFT JOIN manager_messages mm ON u.user_id = mm.user_id
    LEFT JOIN mailing_feedback mf ON u.user_id = mf.user_id
    WHERE u.user_id = ?
    GROUP BY u.user_id
    ''', (user_id,))
    
    user = cursor.fetchone()
    conn.close()
    
    if not user:
        await callback.answer("Пользователь не найден", show_alert=True)
        return
    
    user_name = f"{user['first_name']} {user['last_name'] or ''}".strip()
    username = f"@{user['username']}" if user['username'] else "без username"
    
    feedback = db.get_mailing_feedback_for_user(user_id)
    
    response = f"""
📊 <b>Статистика пользователя</b>

👤 <b>Пользователь:</b> {user_name}
📱 <b>Username:</b> {username}
🆔 <b>ID:</b> {user_id}

<b>Статусы:</b>
• Подписка на рассылку: {'✅ Подписан' if user['mailing_subscribed'] else '❌ Отписан'}
• Заполнил анкету: {'✅ Да' if user['has_filled_questionnaire'] else '❌ Нет'}
• Активен: {'✅ Да' if user['is_active'] else '❌ Нет'}

<b>Активность:</b>
• Анкет: {user['questionnaire_count']}
• Выгрузок: {user['export_count']}
• Сообщений менеджеру: {user['message_count']}
• Отзывов на рассылки: {user['feedback_count']}

<b>Контактные данные:</b>
• Телефон: {user['phone'] or 'Не указан'}
• Email: {user['email'] or 'Не указан'}
• Компания: {user['company'] or 'Не указана'}

<b>Дата регистрации:</b>
{user['created_at'][:19] if user['created_at'] else 'Неизвестно'}
"""
    
    if feedback:
        response += "\n<b>Последние отзывы:</b>\n"
        for i, fb in enumerate(feedback[:3], 1):
            fb_type = "👍" if fb['feedback_type'] == 'like' else "👎" if fb['feedback_type'] == 'dislike' else "💬" if fb['feedback_type'] == 'comment' else "🚫"
            response += f"{i}. {fb_type} {fb['feedback_text'] or fb['feedback_type']} ({fb['created_at'][:16]})\n"
    
    await callback.message.answer(response, parse_mode=ParseMode.HTML)
    await callback.answer()

@dp.callback_query(F.data == "subscription_stats")
async def handle_subscription_stats(callback: types.CallbackQuery):
    """Статистика подписок"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    conn = sqlite3.connect("tenders.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('''
    SELECT 
        COUNT(*) as total_users,
        SUM(CASE WHEN mailing_subscribed = 1 THEN 1 ELSE 0 END) as subscribed,
        SUM(CASE WHEN mailing_subscribed = 0 THEN 1 ELSE 0 END) as unsubscribed,
        SUM(CASE WHEN has_filled_questionnaire = 1 THEN 1 ELSE 0 END) as with_anketa,
        SUM(CASE WHEN has_filled_questionnaire = 0 THEN 1 ELSE 0 END) as without_anketa
    FROM users 
    WHERE is_active = 1
    ''')
    
    stats = cursor.fetchone()
    
    cursor.execute('''
    SELECT COUNT(*) as recent_unsubscribes
    FROM mailing_feedback 
    WHERE feedback_type = 'unsubscribe'
    AND date(created_at) >= date('now', '-30 days')
    ''')
    
    recent = cursor.fetchone()
    
    conn.close()
    
    percentage = (stats['subscribed'] / stats['total_users'] * 100) if stats['total_users'] > 0 else 0
    
    response = f"""
📊 <b>Статистика подписок</b>

<b>Общая статистика:</b>
• Всего активных пользователей: {stats['total_users']}
• Подписано на рассылку: {stats['subscribed']}
• Отписано от рассылки: {stats['unsubscribed']}
• Процент подписки: {percentage:.1f}%

<b>По анкетам:</b>
• С заполненной анкетой: {stats['with_anketa']}
• Без анкеты: {stats['without_anketa']}

<b>Отписки за 30 дней:</b>
• Всего отписок: {recent['recent_unsubscribes']}
"""
    
    await callback.message.answer(response, parse_mode=ParseMode.HTML)
    await callback.answer()

@dp.callback_query(F.data == "refresh_subs")
async def handle_refresh_subs(callback: types.CallbackQuery):
    """Обновление списка подписок"""
    await manage_subscriptions(callback.message)
    await callback.answer("Список обновлен")

@dp.callback_query(F.data.startswith("filter_"))
async def handle_filter_subs(callback: types.CallbackQuery):
    """Фильтрация списка подписок"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    filter_type = callback.data.split("_")[1]
    
    if filter_type == "subscribed":
        users = db.get_users_by_filter("subscribed")
        filter_name = "подписанные"
    elif filter_type == "unsubscribed":
        users = db.get_users_by_filter("unsubscribed")
        filter_name = "отписанные"
    else:
        users = db.get_all_users_with_subscription(30)
        filter_name = "все"
    
    if not users:
        await callback.answer(f"Нет пользователей с фильтром '{filter_name}'", show_alert=True)
        return
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[])
    
    for user in users:
        status_icon = "✅" if user['mailing_subscribed'] else "❌"
        has_anketa = "📋" if user['has_filled_questionnaire'] else "📭"
        
        button_text = f"{status_icon} {has_anketa} {user['first_name']}"
        if user['last_name']:
            button_text += f" {user['last_name']}"
        
        if user['username']:
            button_text += f" (@{user['username']})"
        
        keyboard.inline_keyboard.append([
            InlineKeyboardButton(
                text=button_text[:50],
                callback_data=f"manage_user_{user['user_id']}"
            )
        ])
    
    keyboard.inline_keyboard.append([
        InlineKeyboardButton(text="👥 Все пользователи", callback_data="filter_all"),
        InlineKeyboardButton(text="🔄 Обновить список", callback_data="refresh_subs")
    ])
    
    await callback.message.edit_text(
        f"👥 <b>Управление подписками на рассылку</b>\n\n"
        f"<b>Фильтр:</b> {filter_name}\n"
        f"<b>Найдено пользователей:</b> {len(users)}\n\n"
        "<b>Легенда:</b>\n"
        "✅ - подписан на рассылку\n"
        "❌ - отписан от рассылки\n"
        "📋 - заполнил анкету\n"
        "📭 - без анкеты",
        reply_markup=keyboard,
        parse_mode=ParseMode.HTML
    )
    
    await callback.answer()

# =========== СОЗДАНИЕ РАССЫЛКИ ===========
@dp.message(F.text == "📨 Создать рассылку")
async def start_create_mailing(message: types.Message, state: FSMContext):
    """Начало создания ручной рассылки"""
    if not ADMIN_ID or message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Доступ запрещен", parse_mode=ParseMode.HTML)
        return
    
    await state.set_state(ManualMailing.waiting_for_text)
    await message.answer(
        "📨 <b>Создание ручной рассылки</b>\n\n"
        "Введите текст рассылки. Вы можете использовать HTML-разметку:\n"
        "<b>жирный</b>, <i>курсив</i>, <code>код</code>\n\n"
        "<i>Для отмены нажмите '❌ Отмена'</i>",
        reply_markup=get_cancel_keyboard(),
        parse_mode=ParseMode.HTML
    )

@dp.message(ManualMailing.waiting_for_text)
async def process_mailing_text(message: types.Message, state: FSMContext):
    """Обработка текста рассылки"""
    if message.text == "❌ Отмена":
        await state.clear()
        await message.answer("❌ Создание рассылки отменено.", reply_markup=get_admin_keyboard(), parse_mode=ParseMode.HTML)
        return
    
    await state.update_data(mailing_text=message.text)
    await state.set_state(ManualMailing.waiting_for_filter)
    
    await message.answer(
        "✅ <b>Текст рассылки сохранен</b>\n\n"
        "Теперь выберите категорию пользователей для рассылки:",
        reply_markup=get_mailing_filters_keyboard(),
        parse_mode=ParseMode.HTML
    )

@dp.message(ManualMailing.waiting_for_filter)
async def process_mailing_filter(message: types.Message, state: FSMContext):
    """Обработка фильтра для рассылки"""
    if message.text == "❌ Отмена":
        await state.clear()
        await message.answer("❌ Создание рассылки отменено.", reply_markup=get_admin_keyboard(), parse_mode=ParseMode.HTML)
        return
    
    filter_map = {
        "👥 Все подписанные": "all",
        "📝 С анкетами": "with_questionnaire",
        "📭 Без анкет": "without_questionnaire",
        "🆕 За неделю": "recent_week"
    }
    
    if message.text not in filter_map:
        await message.answer("❌ Пожалуйста, выберите категорию из предложенных кнопок.", parse_mode=ParseMode.HTML)
        return
    
    filter_type = filter_map[message.text]
    
    users = db.get_users_by_filter(filter_type)
    
    if not users:
        await message.answer(
            f"❌ Нет пользователей по выбранному фильтру: {message.text}\n"
            "Попробуйте выбрать другую категорию.",
            reply_markup=get_mailing_filters_keyboard(),
            parse_mode=ParseMode.HTML
        )
        return
    
    await state.update_data(filter_type=filter_type, user_count=len(users))
    await state.set_state(ManualMailing.waiting_for_confirmation)
    
    data = await state.get_data()
    mailing_text = data['mailing_text'][:200] + "..." if len(data['mailing_text']) > 200 else data['mailing_text']
    
    keyboard = ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="✅ Да, отправить")],
            [KeyboardButton(text="❌ Нет, отменить")]
        ],
        resize_keyboard=True
    )
    
    await message.answer(
        f"📨 <b>Подтверждение рассылки</b>\n\n"
        f"<b>Текст:</b>\n{mailing_text}\n\n"
        f"<b>Категория:</b> {message.text}\n"
        f"<b>Количество пользователей:</b> {len(users)}\n\n"
        f"<i>Отправить рассылку?</i>",
        reply_markup=keyboard,
        parse_mode=ParseMode.HTML
    )

@dp.message(ManualMailing.waiting_for_confirmation)
async def process_mailing_confirmation(message: types.Message, state: FSMContext):
    """Подтверждение и отправка рассылки С ОБРАТНОЙ СВЯЗЬЮ"""
    if message.text == "❌ Нет, отменить":
        await state.clear()
        await message.answer("❌ Рассылка отменена.", reply_markup=get_admin_keyboard(), parse_mode=ParseMode.HTML)
        return
    
    if message.text != "✅ Да, отправить":
        await message.answer("❌ Пожалуйста, используйте кнопки для подтверждения.", parse_mode=ParseMode.HTML)
        return
    
    data = await state.get_data()
    mailing_text = data['mailing_text']
    filter_type = data['filter_type']
    user_count = data['user_count']
    
    users = db.get_users_by_filter(filter_type)
    
    if not users:
        await message.answer("❌ Ошибка: пользователи не найдены.", reply_markup=get_admin_keyboard(), parse_mode=ParseMode.HTML)
        await state.clear()
        return
    
    mailing_id = db.create_manual_mailing(
        message.from_user.id,
        mailing_text,
        filter_type,
        json.dumps({"user_count": user_count})
    )
    
    await message.answer(f"🔄 Начинаю отправку рассылки для {len(users)} пользователей...", parse_mode=ParseMode.HTML)
    
    success_count = 0
    failed_count = 0
    
    for user in users:
        try:
            sent_message = await bot.send_message(
                user['user_id'], 
                mailing_text, 
                parse_mode=ParseMode.HTML
            )
            
            sent_message_id = db.save_sent_message(mailing_id, user['user_id'], sent_message.message_id)
            
            feedback_keyboard = get_mailing_feedback_keyboard(sent_message_id)
            await bot.send_message(
                user['user_id'],
                "💬 <b>Как вам эта рассылка?</b>\n\n"
                "Пожалуйста, оставьте обратную связь:",
                reply_markup=feedback_keyboard,
                parse_mode=ParseMode.HTML
            )
            
            success_count += 1
            
            await asyncio.sleep(0.1)
            
        except Exception as e:
            logger.error(f"Не удалось отправить рассылку пользователю {user['user_id']}: {e}")
            failed_count += 1
    
    db.update_mailing_stats(mailing_id, success_count, failed_count)
    
    await message.answer(
        f"✅ <b>Рассылка завершена!</b>\n\n"
        f"📨 <b>ID рассылки:</b> {mailing_id}\n"
        f"👥 <b>Всего пользователей:</b> {len(users)}\n"
        f"✅ <b>Успешно отправлено:</b> {success_count}\n"
        f"❌ <b>Не удалось отправить:</b> {failed_count}\n\n"
        f"<i>Рассылка сохранена в истории. Пользователи получили возможность оставить обратную связь.</i>",
        reply_markup=get_admin_keyboard(),
        parse_mode=ParseMode.HTML
    )
    
    await state.clear()

# =========== ОБРАТНАЯ СВЯЗЬ ПО РАССЫЛКАМ ===========
@dp.callback_query(F.data.startswith("feedback_"))
async def handle_mailing_feedback(callback: types.CallbackQuery, state: FSMContext):
    """Обработка обратной связи по рассылке"""
    try:
        parts = callback.data.split("_")
        feedback_type = parts[1]
        sent_message_id = int(parts[2])
        
        user_id = callback.from_user.id
        username = callback.from_user.username or "без username"
        
        conn = sqlite3.connect("tenders.db")
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT sm.*, mm.mailing_text, mm.id as mailing_id
        FROM sent_messages sm
        JOIN manual_mailings mm ON sm.mailing_id = mm.id
        WHERE sm.id = ? AND sm.user_id = ?
        ''', (sent_message_id, user_id))
        
        sent_message = cursor.fetchone()
        conn.close()
        
        if not sent_message:
            await callback.answer("Сообщение не найдено", show_alert=True)
            return
        
        mailing_id = sent_message['mailing_id']
        
        if feedback_type == "unsubscribe":
            db.toggle_user_mailing_subscription(user_id)
            
            db.save_mailing_feedback(
                mailing_id, 
                user_id, 
                sent_message_id, 
                "unsubscribe", 
                "Пользователь отписался от рассылки"
            )
            
            await callback.message.edit_text(
                callback.message.text + "\n\n✅ <b>Вы отписаны от рассылок</b>",
                reply_markup=None,
                parse_mode=ParseMode.HTML
            )
            
            await callback.answer("Вы отписаны от рассылок")
            
            if ADMIN_ID:
                try:
                    await bot.send_message(
                        ADMIN_ID,
                        f"🚫 <b>ПОЛЬЗОВАТЕЛЬ ОТПИСАЛСЯ ОТ РАССЫЛКИ</b>\n\n"
                        f"👤 Пользователь: @{username}\n"
                        f"🆔 ID: {user_id}\n"
                        f"📨 Рассылка ID: {mailing_id}\n"
                        f"📅 Время: {datetime.now().strftime('%H:%M %d.%m.%Y')}",
                        parse_mode=ParseMode.HTML
                    )
                except Exception as e:
                    logger.error(f"Не удалось уведомить админа об отписке: {e}")
            
            return
        
        elif feedback_type == "comment":
            await state.set_state(FeedbackComment.waiting_for_comment)
            await state.update_data(sent_message_id=sent_message_id, mailing_id=mailing_id)
            
            await callback.message.answer(
                "💬 <b>Напишите ваш комментарий к рассылке:</b>\n\n"
                "<i>Что понравилось или не понравилось? Что можно улучшить?</i>",
                reply_markup=get_cancel_keyboard(),
                parse_mode=ParseMode.HTML
            )
            
            await callback.answer()
            return
        
        else:
            feedback_text_map = {
                "like": "Понравилось",
                "dislike": "Не понравилось"
            }
            
            db.save_mailing_feedback(
                mailing_id, 
                user_id, 
                sent_message_id, 
                feedback_type, 
                feedback_text_map.get(feedback_type, "")
            )
            
            feedback_icon = "👍" if feedback_type == "like" else "👎" if feedback_type == "dislike" else "💬" if feedback_type == "comment" else "🚫"
            await callback.message.edit_text(
                callback.message.text + f"\n\n{feedback_icon} <b>Спасибо за ваш отзыв!</b>",
                reply_markup=None,
                parse_mode=ParseMode.HTML
            )
            
            await callback.answer(f"Спасибо за ваш отзыв: {feedback_text_map.get(feedback_type, '')}")
            
            if ADMIN_ID:
                try:
                    feedback_type_text = "Понравилось" if feedback_type == "like" else "Не понравилось"
                    
                    await bot.send_message(
                        ADMIN_ID,
                        f"{feedback_icon} <b>НОВЫЙ ОТЗЫВ НА РАССЫЛКИ</b>\n\n"
                        f"👤 Пользователь: @{username}\n"
                        f"🆔 ID: {user_id}\n"
                        f"📨 Рассылка ID: {mailing_id}\n"
                        f"💬 Отзыв: {feedback_type_text}\n"
                        f"📅 Время: {datetime.now().strftime('%H:%M %d.%m.%Y')}",
                        parse_mode=ParseMode.HTML
                    )
                except Exception as e:
                    logger.error(f"Не удалось уведомить админа об отзыве: {e}")
    
    except Exception as e:
        logger.error(f"Ошибка обработки обратной связи: {e}")
        await callback.answer("Произошла ошибка", show_alert=True)

@dp.message(FeedbackComment.waiting_for_comment)
async def process_feedback_comment(message: types.Message, state: FSMContext):
    """Обработка комментария к рассылке"""
    if message.text == "❌ Отмена":
        await state.clear()
        await message.answer("❌ Отправка комментария отменена.", reply_markup=get_main_keyboard(), parse_mode=ParseMode.HTML)
        return
    
    data = await state.get_data()
    sent_message_id = data.get('sent_message_id')
    mailing_id = data.get('mailing_id')
    user_id = message.from_user.id
    username = message.from_user.username or "без username"
    
    db.save_mailing_feedback(
        mailing_id, 
        user_id, 
        sent_message_id, 
        "comment", 
        message.text
    )
    
    await message.answer(
        "💬 <b>Спасибо за ваш комментарий!</b>\n\n"
        "Мы учтем ваше мнение для улучшения наших рассылок.",
        reply_markup=get_main_keyboard(),
        parse_mode=ParseMode.HTML
    )
    
    if ADMIN_ID:
        try:
            await bot.send_message(
                ADMIN_ID,
                f"💬 <b>НОВЫЙ КОММЕНТАРИЙ К РАССЫЛКЕ</b>\n\n"
                f"👤 Пользователь: @{username}\n"
                f"🆔 ID: {user_id}\n"
                f"📨 Рассылка ID: {mailing_id}\n"
                f"📝 Комментарий: {message.text[:500]}\n"
                f"📅 Время: {datetime.now().strftime('%H:%M %d.%m.%Y')}",
                parse_mode=ParseMode.HTML
            )
        except Exception as e:
            logger.error(f"Не удалось уведомить админа о комментарии: {e}")
    
    await state.clear()

# =========== ПРОСМОТР ОБРАТНОЙ СВЯЗИ ===========
@dp.message(F.text == "📋 Обратная связь")
async def show_feedback(message: types.Message):
    """Показать обратную связь по рассылкам"""
    if not ADMIN_ID or message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Доступ запрещен", parse_mode=ParseMode.HTML)
        return
    
    conn = sqlite3.connect("tenders.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('''
    SELECT mm.id, mm.mailing_text, mm.created_at, 
           mm.sent_count, mm.feedback_count,
           (SELECT COUNT(DISTINCT mf.user_id) 
            FROM mailing_feedback mf 
            WHERE mf.mailing_id = mm.id) as feedback_users
    FROM manual_mailings mm
    WHERE mm.sent_count > 0
    ORDER BY mm.created_at DESC
    LIMIT 10
    ''')
    
    mailings = cursor.fetchall()
    conn.close()
    
    if not mailings:
        await message.answer("📭 Нет рассылок с обратной связью", parse_mode=ParseMode.HTML)
        return
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[])
    
    for mailing in mailings:
        date_str = mailing['created_at'][:10] if mailing['created_at'] else "??.??.????"
        feedback_percent = (mailing['feedback_count'] / mailing['sent_count'] * 100) if mailing['sent_count'] > 0 else 0
        
        button_text = f"📨 #{mailing['id']} ({date_str}) - {feedback_percent}% отзывов"
        keyboard.inline_keyboard.append([
            InlineKeyboardButton(
                text=button_text,
                callback_data=f"view_feedback_{mailing['id']}"
            )
        ])
    
    keyboard.inline_keyboard.append([
        InlineKeyboardButton(text="📊 Статистика отзывов", callback_data="feedback_stats"),
        InlineKeyboardButton(text="🔄 Обновить", callback_data="refresh_feedback")
    ])
    
    await message.answer(
        "📋 <b>Обратная связь по рассылкам</b>\n\n"
        "Выберите рассылку для просмотра отзывов:",
        reply_markup=keyboard,
        parse_mode=ParseMode.HTML
    )

@dp.callback_query(F.data.startswith("view_feedback_"))
async def handle_view_feedback(callback: types.CallbackQuery):
    """Просмотр обратной связи по конкретной рассылке"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    mailing_id = int(callback.data.split("_")[2])
    
    feedback = db.get_mailing_feedback(mailing_id)
    
    if not feedback:
        await callback.answer("Нет обратной связи по этой рассылке", show_alert=True)
        return
    
    likes = sum(1 for f in feedback if f['feedback_type'] == 'like')
    dislikes = sum(1 for f in feedback if f['feedback_type'] == 'dislike')
    comments = sum(1 for f in feedback if f['feedback_type'] == 'comment')
    unsubscribes = sum(1 for f in feedback if f['feedback_type'] == 'unsubscribe')
    
    response = f"""
📋 <b>Обратная связь по рассылке #{mailing_id}</b>

<b>Статистика:</b>
• Всего отзывов: {len(feedback)}
• 👍 Понравилось: {likes}
• 👎 Не понравилось: {dislikes}
• 💬 Комментарии: {comments}
• 🚫 Отписки: {unsubscribes}

<b>Последние отзывы:</b>
"""
    
    for i, fb in enumerate(feedback[:10], 1):
        fb_type = "👍" if fb['feedback_type'] == 'like' else "👎" if fb['feedback_type'] == 'dislike' else "💬" if fb['feedback_type'] == 'comment' else "🚫"
        user_name = f"@{fb['username']}" if fb['username'] else f"{fb['first_name']} {fb['last_name'] or ''}"
        date_str = fb['created_at'][:16] if fb['created_at'] else "??.?? ??:??"
        
        response += f"\n{i}. {fb_type} <b>{user_name}</b> ({date_str})"
        if fb['feedback_text']:
            response += f"\n   {fb['feedback_text'][:100]}..."
    
    if len(feedback) > 10:
        response += f"\n\n... и еще {len(feedback) - 10} отзывов"
    
    await callback.message.answer(response, parse_mode=ParseMode.HTML)
    await callback.answer()

@dp.callback_query(F.data == "feedback_stats")
async def handle_feedback_stats(callback: types.CallbackQuery):
    """Статистика обратной связи"""
    if not ADMIN_ID or callback.from_user.id != ADMIN_ID:
        await callback.answer("⛔ Доступ запрещен", show_alert=True)
        return
    
    conn = sqlite3.connect("tenders.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('''
    SELECT 
        COUNT(*) as total_feedback,
        SUM(CASE WHEN feedback_type = 'like' THEN 1 ELSE 0 END) as likes,
        SUM(CASE WHEN feedback_type = 'dislike' THEN 1 ELSE 0 END) as dislikes,
        SUM(CASE WHEN feedback_type = 'comment' THEN 1 ELSE 0 END) as comments,
        SUM(CASE WHEN feedback_type = 'unsubscribe' THEN 1 ELSE 0 END) as unsubscribes
    FROM mailing_feedback
    ''')
    
    stats = cursor.fetchone()
    
    cursor.execute('''
    SELECT 
        COUNT(*) as recent_feedback,
        SUM(CASE WHEN feedback_type = 'unsubscribe' THEN 1 ELSE 0 END) as recent_unsubscribes
    FROM mailing_feedback 
    WHERE date(created_at) >= date('now', '-30 days')
    ''')
    
    recent = cursor.fetchone()
    
    cursor.execute('''
    SELECT mm.id, mm.mailing_text, COUNT(mf.id) as feedback_count
    FROM manual_mailings mm
    LEFT JOIN mailing_feedback mf ON mm.id = mf.mailing_id
    GROUP BY mm.id
    ORDER BY feedback_count DESC
    LIMIT 5
    ''')
    
    popular = cursor.fetchall()
    
    conn.close()
    
    response = f"""
📊 <b>Статистика обратной связи</b>

<b>Общая статистика:</b>
• Всего отзывов: {stats['total_feedback'] or 0}
• 👍 Понравилось: {stats['likes'] or 0}
• 👎 Не понравилось: {stats['dislikes'] or 0}
• 💬 Комментарии: {stats['comments'] or 0}
• 🚫 Отписки: {stats['unsubscribes'] or 0}

<b>За последние 30 дней:</b>
• Новых отзывов: {recent['recent_feedback'] or 0}
• Отписок: {recent['recent_unsubscribes'] or 0}

<b>Самые обсуждаемые рассылки:</b>
"""
    
    for i, mailing in enumerate(popular, 1):
        mailing_text_preview = mailing['mailing_text'][:50] + "..." if len(mailing['mailing_text']) > 50 else mailing['mailing_text']
        response += f"\n{i}. ID#{mailing['id']}: {mailing_text_preview}"
        response += f"\n   Отзывов: {mailing['feedback_count']}"
    
    await callback.message.answer(response, parse_mode=ParseMode.HTML)
    await callback.answer()

@dp.callback_query(F.data == "refresh_feedback")
async def handle_refresh_feedback(callback: types.CallbackQuery):
    """Обновление списка обратной связи"""
    await show_feedback(callback.message)
    await callback.answer("Список обновлен")

# =========== ОСТАЛЬНЫЕ АДМИН ФУНКЦИИ ===========
@dp.message(F.text == "📩 Сообщения менеджеру")
async def show_manager_messages(message: types.Message):
    """Показать сообщения менеджеру"""
    if not ADMIN_ID or message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Доступ запрещен", parse_mode=ParseMode.HTML)
        return
    
    conn = sqlite3.connect("tenders.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('''
    SELECT mm.*, u.username, u.first_name, u.last_name 
    FROM manager_messages mm
    JOIN users u ON mm.user_id = u.user_id
    WHERE mm.processed = 0
    ORDER BY mm.created_at DESC
    LIMIT 10
    ''')
    
    messages = cursor.fetchall()
    conn.close()
    
    if not messages:
        await message.answer("📭 Новых сообщений менеджеру нет", parse_mode=ParseMode.HTML)
        return
    
    response = f"📩 <b>Новые сообщения менеджеру ({len(messages)}):</b>\n\n"
    
    for i, msg in enumerate(messages, 1):
        date_str = msg['created_at'][:16] if msg['created_at'] else "??.?? ??:??"
        type_icon = "💬" if msg['message_type'] == 'text' else "📎" if msg['message_type'] == 'document' else "🖼"
        
        response += f"{i}. <b>#{msg['id']}</b> {type_icon}\n"
        response += f"   👤 @{msg['username'] or 'без username'}\n"
        response += f"   📝 {msg['message_text'][:50]}...\n"
        response += f"   ⏰ {date_str}\n\n"
    
    await message.answer(response, parse_mode=ParseMode.HTML)

@dp.message(F.text == "⚙️ Настройки")
async def show_settings(message: types.Message):
    """Показать настройки"""
    if not ADMIN_ID or message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Доступ запрещен", parse_mode=ParseMode.HTML)
        return
    
    stats = db.get_statistics(7)
    
    await message.answer(
        "⚙️ <b>Настройки бота:</b>\n\n"
        "<b>Текущие параметры:</b>\n"
        f"• Время работы: {WORK_START_HOUR}:00-{WORK_END_HOUR}:00 Пн-Пт\n"
        f"• Follow-up через: 1 час\n"
        f"• ID администратора: {ADMIN_ID}\n\n"
        "<b>Статистика за неделю:</b>\n"
        f"• Новых пользователей: {stats['new_users']}\n"
        f"• Пользователей с подпиской: {stats['subscribed_users']}\n"
        f"• Пользователей без подписки: {stats['unsubscribed_users']}\n"
        f"• Получено отзывов: {stats['mailings_feedback']}\n\n"
        "<b>Функции:</b>\n"
        "✅ Отправка анкет в Word\n"
        "✅ Диалог с менеджером\n"
        "✅ Ручные рассылки с обратной связью\n"
        "✅ Управление подписками\n"
        "✅ Просмотр обратной связи\n"
        "✅ Автоматические отчеты\n\n"
        "<i>Для изменения настроек обратитесь к разработчику</i>",
        parse_mode=ParseMode.HTML
    )

@dp.message(F.text == "👤 Режим пользователя")
async def switch_to_user_mode(message: types.Message, state: FSMContext):
    """Переключение в режим пользователя"""
    if not ADMIN_ID or message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Доступ запрещен", parse_mode=ParseMode.HTML)
        return
    
    await state.clear()
    await message.answer(
        "👤 <b>Вы перешли в режим пользователя</b>\n\n"
        "Теперь вы можете тестировать функции бота как обычный пользователь.\n\n"
        "Чтобы вернуться в панель администратора, используйте команду /admin",
        reply_markup=get_main_keyboard(),
        parse_mode=ParseMode.HTML
    )

# =========== ЗАПОЛНЕНИЕ АНКЕТЫ (ПЕРВАЯ ЧАСТЬ - 1-4 ВОПРОСЫ) ===========
@dp.message(Questionnaire.waiting_for_activity)
async def process_activity(message: types.Message, state: FSMContext):
    """Обработка сферы деятельности (первый вопрос)"""
    await state.update_data(activity=message.text.strip())
    await message.answer(
        "✅ <b>Сфера деятельности сохранена</b>\n\n"
        "Введите <b>регионы работы</b> (города, области):\n"
        "<i>Пример: Москва, Московская область, Санкт-Петербург</i>",
        parse_mode=ParseMode.HTML
    )
    await state.set_state(Questionnaire.waiting_for_region)

@dp.message(Questionnaire.waiting_for_region)
async def process_region(message: types.Message, state: FSMContext):
    """Обработка регионов"""
    await state.update_data(region=message.text.strip())
    await message.answer(
        "✅ <b>Регионы сохранены</b>\n\n"
        "Укажите <b>предпочтительный бюджет контрактов</b>:\n"
        "<i>Пример: от 100 000 до 1 000 000 руб.</i>",
        parse_mode=ParseMode.HTML
    )
    await state.set_state(Questionnaire.waiting_for_budget)

@dp.message(Questionnaire.waiting_for_budget)
async def process_budget(message: types.Message, state: FSMContext):
    """Обработка бюджета"""
    await state.update_data(budget=message.text.strip())
    await message.answer(
        "✅ <b>Бюджет сохранен</b>\n\n"
        "Введите <b>ключевые слова для поиска</b> (через запятую):\n"
        "<i>Пример: строительные работы, поставка оборудования, IT-аутсорсинг</i>",
        parse_mode=ParseMode.HTML
    )
    await state.set_state(Questionnaire.waiting_for_keywords)

@dp.message(Questionnaire.waiting_for_keywords)
async def process_keywords(message: types.Message, state: FSMContext):
    """Обработка ключевых слов и завершение первой части анкеты"""
    user_data = await state.get_data()
    user_data['keywords'] = message.text.strip()
    
    user_id = message.from_user.id
    username = message.from_user.username or "без username"
    
    # Сохраняем первую часть анкеты (только 1-4 вопросы)
    questionnaire_id = db.save_questionnaire_partial(user_id, user_data)
    
    # Отправляем админу уведомление о первой части анкеты
    await send_partial_questionnaire_to_admin(questionnaire_id, user_id, user_data, username)
    
    await message.answer(
        "🎉 <b>Анкета отправлена менеджеру для анализа!</b>\n\n"
        "✅ <b>Ваши данные (1-4 пункты) сохранены:</b>\n"
        f"• Сфера: {user_data['activity'][:50]}\n"
        f"• Регионы: {user_data['region'][:50]}\n"
        f"• Бюджет: {user_data['budget'][:50]}\n"
        f"• Ключевые слова: {user_data['keywords'][:50]}\n\n"
        "<i>Мы проанализируем вашу анкету и вернемся с выгрузкой тендеров.</i>\n"
        "<i>Обычно это занимает от 1 до 24 часов в рабочее время.</i>",
        reply_markup=get_main_keyboard(),
        parse_mode=ParseMode.HTML
    )
    
    await state.clear()

# =========== ОБРАБОТЧИК ЗАПОЛНЕНИЯ КОНТАКТОВ ДЛЯ ВЫГРУЗКИ ===========
@dp.callback_query(F.data.startswith("fill_contacts_"))
async def handle_fill_contacts(callback: types.CallbackQuery, state: FSMContext):
    """Начало заполнения контактов для получения выгрузки"""
    export_id = int(callback.data.split("_")[2])
    
    await state.update_data(export_id=export_id)
    await state.set_state(ExportContacts.waiting_for_company)
    
    await callback.message.edit_text(
        callback.message.text + "\n\n📝 <b>Заполните оставшиеся данные:</b>",
        parse_mode=ParseMode.HTML
    )
    
    await callback.message.answer(
        "🏢 Введите <b>полное название вашей компании</b>:",
        reply_markup=get_cancel_keyboard(),
        parse_mode=ParseMode.HTML
    )
    
    await callback.answer()

@dp.message(ExportContacts.waiting_for_company)
async def process_export_company(message: types.Message, state: FSMContext):
    """Обработка названия компании для выгрузки"""
    await state.update_data(company_name=message.text.strip())
    await message.answer(
        "✅ <b>Компания сохранена</b>\n\n"
        "Введите ваше <b>ФИО полностью</b>:",
        parse_mode=ParseMode.HTML
    )
    await state.set_state(ExportContacts.waiting_for_name)

@dp.message(ExportContacts.waiting_for_name)
async def process_export_name(message: types.Message, state: FSMContext):
    """Обработка ФИО для выгрузки"""
    await state.update_data(full_name=message.text.strip())
    await message.answer(
        "✅ <b>ФИО сохранено</b>\n\n"
        "Теперь укажите ваш <b>телефон для связи</b>.\n\n"
        "Вы можете нажать кнопку ниже, чтобы поделиться телефоном, или ввести его вручную:",
        reply_markup=get_phone_keyboard_simple(),
        parse_mode=ParseMode.HTML
    )
    await state.set_state(ExportContacts.waiting_for_phone)

@dp.message(ExportContacts.waiting_for_phone)
async def process_export_phone(message: types.Message, state: FSMContext):
    """Обработка телефона для выгрузки с кнопкой поделиться"""
    if message.text == "❌ Отмена":
        await state.clear()
        await message.answer(
            "❌ Заполнение контактов отменено.",
            reply_markup=get_main_keyboard(),
            parse_mode=ParseMode.HTML
        )
        return
    
    phone = None
    
    if message.contact:
        phone = message.contact.phone_number
    elif message.text and message.text != "📱 Поделиться телефоном":
        phone = message.text.strip()
    else:
        await message.answer(
            "❌ Пожалуйста, укажите ваш телефон для связи.\n\n"
            "Вы можете нажать кнопку '📱 Поделиться телефоном' или ввести номер вручную.",
            reply_markup=get_phone_keyboard_simple(),
            parse_mode=ParseMode.HTML
        )
        return
    
    await state.update_data(phone=phone)
    await message.answer(
        f"✅ <b>Телефон сохранен: {phone}</b>\n\n"
        "Введите ваш <b>email для отправки тендеров</b>:",
        reply_markup=get_cancel_keyboard(),
        parse_mode=ParseMode.HTML
    )
    await state.set_state(ExportContacts.waiting_for_email)

@dp.message(ExportContacts.waiting_for_email)
async def process_export_email(message: types.Message, state: FSMContext):
    """Завершение заполнения контактов и отправка выгрузки"""
    if message.text == "❌ Отмена":
        await state.clear()
        await message.answer(
            "❌ Заполнение контактов отменено.",
            reply_markup=get_main_keyboard(),
            parse_mode=ParseMode.HTML
        )
        return
    
    data = await state.get_data()
    data['email'] = message.text.strip()
    
    export_id = data.get('export_id')
    user_id = message.from_user.id
    
    # Сохраняем полную анкету (обновляем частичную)
    db.update_partial_to_complete(user_id, data)
    
    # Получаем данные о выгрузке
    export = db.get_export_by_id(export_id)
    
    if export:
        if export['file_path'] and os.path.exists(export['file_path']):
            # Отправляем файл пользователю
            await send_export_file_to_user(user_id, export['file_path'], export['file_name'], export_id)
            
            await message.answer(
                "🎉 <b>Спасибо! Выгрузка тендеров отправлена!</b>\n\n"
                "✅ <b>Ваши контакты сохранены:</b>\n"
                f"• Компания: {data.get('company_name')}\n"
                f"• ФИО: {data.get('full_name')}\n"
                f"• Телефон: {data.get('phone')}\n"
                f"• Email: {data.get('email')}\n\n"
                "<i>Файл с выгрузкой тендеров отправлен вам выше.</i>\n"
                "<i>Вы также можете найти его в разделе '📊 Мои выгрузки'</i>",
                reply_markup=get_main_keyboard(),
                parse_mode=ParseMode.HTML
            )
        else:
            # Если файла нет, отправляем уведомление
            await bot.send_message(
                user_id,
                f"📨 <b>Ваша выгрузка тендеров #{export_id} готова!</b>\n\n"
                f"✅ <b>Ваши контакты сохранены:</b>\n"
                f"• Компания: {data.get('company_name')}\n"
                f"• ФИО: {data.get('full_name')}\n"
                f"• Телефон: {data.get('phone')}\n"
                f"• Email: {data.get('email')}\n\n"
                "<i>Выгрузка была успешно подготовлена. "
                f"Вы можете посмотреть ее в разделе '📊 Мои выгрузки'.</i>",
                parse_mode=ParseMode.HTML
            )
            
            # Обновляем статус выгрузки
            db.mark_export_completed(export_id, "Автоматическая отправка")
        
        # Отмечаем запрос контактов как выполненный
        db.mark_contact_request_completed(export_id)
        
        # Уведомляем админа
        if ADMIN_ID:
            try:
                await bot.send_message(
                    ADMIN_ID,
                    f"✅ <b>Пользователь заполнил контакты и получил выгрузку</b>\n\n"
                    f"👤 Пользователь ID: {user_id}\n"
                    f"📱 Username: @{message.from_user.username or 'без username'}\n"
                    f"🏢 Компания: {data.get('company_name')}\n"
                    f"📞 Телефон: {data.get('phone')}\n"
                    f"📧 Email: {data.get('email')}\n"
                    f"📄 Выгрузка ID: {export_id}\n\n"
                    f"<i>Выгрузка отправлена пользователю.</i>",
                    parse_mode=ParseMode.HTML
                )
            except Exception as e:
                logger.error(f"Не удалось уведомить админа: {e}")
    else:
        await message.answer(
            "✅ <b>Ваши контакты сохранены!</b>\n\n"
            "<i>Выгрузка будет отправлена вам в ближайшее время менеджером.</i>",
            reply_markup=get_main_keyboard(),
            parse_mode=ParseMode.HTML
        )
    
    await state.clear()

# =========== ЗАПУСК БОТА И HTTP СЕРВЕРА ===========
async def start_http_server():
    """Запуск HTTP сервера для Railway"""
    app = web.Application()
    app.router.add_get('/', health_check)
    app.router.add_get('/status', status_check)
    
    runner = web.AppRunner(app)
    await runner.setup()
    
    site = web.TCPSite(runner, '0.0.0.0', PORT)
    await site.start()
    
    print(f"✅ HTTP сервер запущен на порту {PORT}")
    print(f"✅ Health check доступен по пути: /")
    print(f"✅ Статус бота: /status")
    
    return runner

async def main():
    """Основная функция запуска"""
    print("\n" + "="*60)
    print("🚀 ЗАПУСК БОТА ТРИТИКА (ТЕНДЕРПОИСК)")
    print("="*60)
    
    # Создаем папку для экспортов если её нет
    os.makedirs(EXPORTS_DIR, exist_ok=True)
    print(f"✅ Папка для выгрузок создана: {EXPORTS_DIR}")
    
    # Скачиваем файл анкеты при запуске
    print("📥 Проверяю наличие файла анкеты...")
    if not os.path.exists(ANKETA_LOCAL_PATH):
        print("Файл анкеты не найден локально, скачиваю с GitHub...")
        success = await download_anketa_file()
        if not success:
            print("⚠️ Внимание: Файл анкеты не скачан. Будет использоваться ссылка на GitHub.")
    else:
        file_size = os.path.getsize(ANKETA_LOCAL_PATH)
        print(f"✅ Файл анкеты уже существует ({file_size} байт)")
    
    # Проверяем бота
    try:
        bot_info = await bot.get_me()
        print(f"✅ Бот: @{bot_info.username}")
        print(f"✅ Имя: {bot_info.first_name}")
        print(f"✅ ID: {bot_info.id}")
        
        if ADMIN_ID:
            print(f"✅ Администратор: {ADMIN_ID}")
        else:
            print("⚠️ Администратор не установлен (ADMIN_ID)")
    except Exception as e:
        print(f"❌ Ошибка проверки бота: {e}")
        print("⚠️ Проверьте токен бота")
        return
    
    # Запускаем HTTP сервер
    try:
        http_runner = await start_http_server()
    except Exception as e:
        print(f"❌ Ошибка запуска HTTP сервера: {e}")
        print("⚠️ Возможно, порт {PORT} уже занят")
        return
    
    # Запускаем задачу для follow-up сообщений
    asyncio.create_task(schedule_follow_ups())
    print("✅ Follow-up система запущена")
    
    # Очищаем вебхуки и добавляем небольшую задержку
    await bot.delete_webhook(drop_pending_updates=True)
    await asyncio.sleep(1)
    print("✅ Вебхуки очищены")
    
    print("\n" + "="*60)
    print("🤖 БОТ УСПЕШНО ЗАПУЩЕН!")
    print("="*60)
    print(f"\n📱 Откройте Telegram и найдите бота:")
    print(f"   👉 https://t.me/{bot_info.username}")
    print("\n👤 Обычный режим: /start")
    print("🛠️ Админ-панель: /admin (если настроен ADMIN_ID)")
    print("\n🔄 Ожидание сообщений...")
    print(f"🌐 Health check активен на порту {PORT}\n")
    print("⏰ Follow-up система активна (проверка каждые 5 минут)")
    print("📨 Система запроса контактов для выгрузок активна")
    print("📱 Кнопка 'Поделиться телефоном' добавлена во вторую часть анкеты (при запросе контактов)")
    print("📱 Кнопка 'Поделиться телефоном' удалена из главного меню")
    
    # Запускаем polling бота
    try:
        await dp.start_polling(bot, skip_updates=True)
    except KeyboardInterrupt:
        print("\n🛑 Бот остановлен пользователем")
    except Exception as e:
        logger.error(f"Ошибка в работе бота: {e}")
        print(f"❌ Ошибка: {e}")
    finally:
        # Очищаем ресурсы
        await http_runner.cleanup()
        await bot.session.close()
        print("👋 Сессия бота закрыта")

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n🛑 Приложение остановлено")
    except Exception as e:
        logger.error(f"Критическая ошибка при запуске: {e}")
        print(f"❌ Критическая ошибка: {e}")
